from __future__ import annotations

import json
import sys
import time
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import AsyncMock, patch

from docx import Document


sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from sub2api_worker import main as worker  # noqa: E402


class AcademicPaperWorkflowTests(unittest.IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        worker.canceled_runs.clear()
        worker.active_model_proxy_tasks.clear()

    def tearDown(self) -> None:
        worker.canceled_runs.clear()
        worker.active_model_proxy_tasks.clear()

    @staticmethod
    def payload(
        *,
        input_values: dict[str, object] | None = None,
        artifacts: list[worker.WorkerArtifactRef] | None = None,
    ) -> worker.WorkerRunRequest:
        payload = worker.WorkerRunRequest(
            run_id=901,
            app_id=91,
            app_version_id=92,
            run_token="paper-secret",
            callback_url="https://sub2api.test/api/v1/agent-runs/901/callback",
            model_proxy_url="https://sub2api.test/api/v1/agent-runs/901/model-proxy",
            artifact_url="https://sub2api.test/api/v1/agent-runs/901/artifacts",
            user={"user_id": 31, "api_key_id": 32},
            input=input_values or {},
            input_artifacts=artifacts or [],
            metadata={"worker_route": "/academic-paper/runs"},
        )
        payload.node_model_policy = {
            "academic_paper.plan": worker.ModelPolicy(
                node_id="academic_paper",
                role="plan",
                model="gpt-5.5",
                capability="text",
            ),
            "academic_paper.write": worker.ModelPolicy(
                node_id="academic_paper",
                role="write",
                model="gpt-5.5",
                capability="text",
            ),
        }
        return payload

    async def test_health_and_route_expose_academic_paper(self) -> None:
        health = await worker.health()
        self.assertIn("academic_paper", health["capabilities"])
        self.assertIn("/academic-paper/runs", health["routes"]["runs"])
        paths = {route.path for route in worker.app.routes}
        self.assertIn("/academic-paper/runs", paths)
        self.assertTrue(worker.is_academic_paper_run(self.payload()))

    async def test_academic_paper_model_proxy_retries_transient_502_then_succeeds(self) -> None:
        payload = self.payload()
        policy = worker.SelectedPolicy(
            policy_key="academic_paper.write",
            node_id="academic_paper",
            role="write",
            model="gpt-5.5",
            capability="text",
        )
        transient = worker.WorkerFailure(
            "MODEL_PROXY_FAILED",
            json.dumps(
                {
                    "code": 502,
                    "message": "Upstream service temporarily unavailable",
                    "reason": "AGENT_MODEL_PROXY_UPSTREAM_ERROR",
                }
            ),
        )
        success = {"response": {"text": "ok"}, "usage": {"total_tokens": 10}}
        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=[transient, success])) as proxy,
            patch.object(worker, "PAPER_MODEL_PROXY_MAX_ATTEMPTS", 3),
            patch.object(worker, "PAPER_MODEL_PROXY_RETRY_BASE_SECONDS", 0),
            patch.object(worker, "PAPER_MODEL_PROXY_RETRY_MAX_SECONDS", 0),
        ):
            result = await worker.call_academic_paper_model_proxy(payload, policy, "prompt", stage="correction")

        self.assertEqual(success, result)
        self.assertEqual(2, proxy.await_count)

    async def test_academic_paper_model_proxy_retries_generic_stream_upstream_failure(self) -> None:
        payload = self.payload()
        policy = worker.SelectedPolicy(
            policy_key="academic_paper.write",
            node_id="academic_paper",
            role="write",
            model="gpt-5.5",
            capability="text",
        )
        transient = worker.WorkerFailure(
            "MODEL_PROXY_STREAM_FAILED",
            "Upstream request failed",
        )
        success = {"response": {"text": "ok"}, "usage": {"total_tokens": 10}}
        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=[transient, success])) as proxy,
            patch.object(worker, "PAPER_MODEL_PROXY_MAX_ATTEMPTS", 3),
            patch.object(worker, "PAPER_MODEL_PROXY_RETRY_BASE_SECONDS", 0),
            patch.object(worker, "PAPER_MODEL_PROXY_RETRY_MAX_SECONDS", 0),
        ):
            result = await worker.call_academic_paper_model_proxy(payload, policy, "prompt", stage="correction")

        self.assertEqual(success, result)
        self.assertEqual(2, proxy.await_count)

    async def test_academic_paper_model_proxy_exhausts_transient_retries(self) -> None:
        payload = self.payload()
        policy = worker.SelectedPolicy(
            policy_key="academic_paper.write",
            node_id="academic_paper",
            role="write",
            model="gpt-5.5",
            capability="text",
        )
        failures = [
            worker.WorkerFailure(
                "MODEL_PROXY_FAILED",
                json.dumps(
                    {
                        "code": 502,
                        "message": "Upstream service temporarily unavailable",
                        "reason": "AGENT_MODEL_PROXY_UPSTREAM_ERROR",
                    }
                ),
            )
            for _ in range(3)
        ]
        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=failures)) as proxy,
            patch.object(worker, "PAPER_MODEL_PROXY_MAX_ATTEMPTS", 3),
            patch.object(worker, "PAPER_MODEL_PROXY_RETRY_BASE_SECONDS", 0),
            patch.object(worker, "PAPER_MODEL_PROXY_RETRY_MAX_SECONDS", 0),
        ):
            with self.assertRaises(worker.WorkerFailure) as raised:
                await worker.call_academic_paper_model_proxy(payload, policy, "prompt", stage="review")

        self.assertEqual("MODEL_PROXY_FAILED", raised.exception.code)
        self.assertEqual(3, proxy.await_count)

    async def test_academic_paper_model_proxy_does_not_retry_non_transient_error(self) -> None:
        payload = self.payload()
        policy = worker.SelectedPolicy(
            policy_key="academic_paper.plan",
            node_id="academic_paper",
            role="plan",
            model="gpt-5.5",
            capability="text",
        )
        non_transient = worker.WorkerFailure(
            "MODEL_PROXY_FAILED",
            json.dumps(
                {
                    "code": 401,
                    "message": "API key is invalid",
                    "reason": "AGENT_MODEL_PROXY_API_KEY_UNAVAILABLE",
                }
            ),
        )
        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=non_transient)) as proxy,
            patch.object(worker, "PAPER_MODEL_PROXY_MAX_ATTEMPTS", 3),
            patch.object(worker, "PAPER_MODEL_PROXY_RETRY_BASE_SECONDS", 0),
        ):
            with self.assertRaises(worker.WorkerFailure) as raised:
                await worker.call_academic_paper_model_proxy(payload, policy, "prompt", stage="plan")

        self.assertIs(non_transient, raised.exception)
        self.assertEqual(1, proxy.await_count)

    async def test_academic_paper_model_proxy_translates_cancellation_during_backoff(self) -> None:
        payload = self.payload()
        policy = worker.SelectedPolicy(
            policy_key="academic_paper.write",
            node_id="academic_paper",
            role="write",
            model="gpt-5.5",
            capability="text",
        )
        transient = worker.WorkerFailure(
            "MODEL_PROXY_FAILED",
            json.dumps(
                {
                    "code": 502,
                    "message": "Upstream service temporarily unavailable",
                    "reason": "AGENT_MODEL_PROXY_UPSTREAM_ERROR",
                }
            ),
        )

        async def cancel_during_backoff(_: float) -> None:
            worker.canceled_runs.add(payload.run_id)
            raise worker.asyncio.CancelledError()

        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=transient)) as proxy,
            patch.object(worker, "PAPER_MODEL_PROXY_MAX_ATTEMPTS", 3),
            patch.object(worker, "PAPER_MODEL_PROXY_RETRY_BASE_SECONDS", 0.01),
            patch.object(worker.asyncio, "sleep", new=cancel_during_backoff),
        ):
            with self.assertRaises(worker.WorkerCanceled):
                await worker.call_academic_paper_model_proxy(payload, policy, "prompt", stage="correction")

        self.assertEqual(1, proxy.await_count)

    def test_academic_paper_model_proxy_classifies_429_and_network_errors_as_transient(self) -> None:
        rate_limited = worker.WorkerFailure(
            "MODEL_PROXY_FAILED",
            json.dumps({"code": 429, "message": "Too many requests", "reason": "RATE_LIMITED"}),
        )
        generic_upstream_failure = worker.WorkerFailure(
            "MODEL_PROXY_STREAM_FAILED",
            "Upstream request failed",
        )
        timeout = worker.httpx.ReadTimeout(
            "read timed out",
            request=worker.httpx.Request("POST", "https://sub2api.test/model-proxy"),
        )
        self.assertTrue(worker.is_transient_academic_paper_model_error(rate_limited))
        self.assertTrue(worker.is_transient_academic_paper_model_error(generic_upstream_failure))
        self.assertTrue(worker.is_transient_academic_paper_model_error(timeout))

    def test_plan_parser_accepts_markdown_fence_and_has_invalid_json_fallback(self) -> None:
        raw = """```json
        {"title":"模型题目","abstract":"摘要","keywords":["人工智能"],"sections":[
          {"title":"绪论","target_words":200},
          {"title":"结论","target_words":300}
        ]}
        ```"""
        plan = worker.parse_academic_paper_plan(
            raw,
            {"topic": "人工智能教育", "paper_title": "用户指定题目"},
            1000,
        )
        self.assertEqual("用户指定题目", plan["title"])
        self.assertEqual(1000, sum(item["target_words"] for item in plan["sections"]))

        fallback = worker.parse_academic_paper_plan(
            "这不是 JSON",
            {"topic": "数字化转型", "outline_requirements": "第一章 绪论\n第二章 结论"},
            2000,
        )
        self.assertEqual("数字化转型", fallback["title"])
        self.assertGreaterEqual(len(fallback["sections"]), 2)
        self.assertEqual(2000, sum(item["target_words"] for item in fallback["sections"]))

    def test_academic_paper_keywords_obey_enabled_count_contract(self) -> None:
        values = {
            "topic": "短视频内容营销对大学生消费意愿的影响",
            "discipline": "市场营销",
            "research_method": "问卷调查、文献分析",
            "keywords_enabled": True,
            "keywords_count": "5",
        }
        plan = worker.parse_academic_paper_plan(
            json.dumps(
                {
                    "title": "短视频内容营销研究",
                    "abstract": "摘要",
                    "keywords": ["短视频平台", "内容营销", "大学生", "消费意愿", "影响机制", "问卷调查"],
                    "sections": [{"title": "绪论"}, {"title": "结论"}],
                },
                ensure_ascii=False,
            ),
            values,
            1000,
        )
        self.assertEqual(["短视频平台", "内容营销", "大学生", "消费意愿", "影响机制"], plan["keywords"])

        worker.apply_academic_paper_review(
            plan,
            [],
            {
                "keywords": ["内容营销", "大学生", "消费意愿", "短视频平台", "购买行为", "平台机制"],
            },
            values,
        )
        self.assertEqual(["内容营销", "大学生", "消费意愿", "短视频平台", "购买行为"], plan["keywords"])

        paper = worker.build_academic_paper_payload(plan, [], values)
        self.assertEqual(5, len(paper["keywords"]))
        self.assertEqual(5, len(set(paper["keywords"])))
        self.assertIn("恰好返回 5 个", worker.build_academic_paper_plan_prompt(values, 1000, ""))
        self.assertIn("恰好返回 5 个", worker.build_academic_paper_review_prompt(values, plan, []))

        sparse = worker.normalize_academic_paper_keywords(["数字治理", "数字治理"], values)
        self.assertEqual(5, len(sparse))
        self.assertEqual(5, len(set(sparse)))
        self.assertEqual(2, len(worker.normalize_academic_paper_keywords([], {**values, "keywords_count": 1})))
        self.assertEqual(10, len(worker.normalize_academic_paper_keywords([], {**values, "keywords_count": 99})))
        self.assertEqual([], worker.normalize_academic_paper_keywords(["不应保留"], {**values, "keywords_enabled": False}))

    def test_locked_outline_spec_is_validated_allocated_and_built_as_fixed_tree(self) -> None:
        nodes = worker.parse_academic_paper_outline_spec(
            {
                "version": 1,
                "nodes": [
                    {"id": "chapter-1", "title": "绪论", "level": 1},
                    {"id": "section-1-1", "title": "研究背景", "level": 2},
                    {"id": "section-1-1-1", "title": "现实背景", "level": 3},
                    {"id": "chapter-2", "title": "结论", "level": 1},
                ],
            },
            1000,
        )
        self.assertEqual(
            1000 - worker.academic_paper_abstract_word_budget(1000),
            sum(node["target_words"] for node in nodes),
        )
        self.assertEqual(["chapter-1", "section-1-1", "section-1-1-1", "chapter-2"], [node["id"] for node in nodes])

        nodes_without_abstract = worker.parse_academic_paper_outline_spec(
            {
                "version": 1,
                "nodes": [{"id": "chapter-1", "title": "正文", "level": 1}],
            },
            1000,
            abstract_enabled=False,
        )
        self.assertEqual(1000, sum(node["target_words"] for node in nodes_without_abstract))

        tree = worker.build_locked_academic_paper_tree(
            nodes,
            {node["id"]: f"{node['id']} 正文" for node in nodes},
        )
        self.assertEqual(["绪论", "结论"], [section["title"] for section in tree])
        self.assertEqual("研究背景", tree[0]["children"][0]["title"])
        self.assertEqual("现实背景", tree[0]["children"][0]["children"][0]["title"])
        self.assertTrue(worker.locked_academic_paper_outline_matches(tree, nodes))

    def test_locked_paper_word_count_limits_are_85_to_115_percent(self) -> None:
        self.assertEqual((85, 115), worker.locked_paper_word_count_limits(100))
        self.assertTrue(worker.paper_word_count_within_tolerance(85, 100))
        self.assertTrue(worker.paper_word_count_within_tolerance(115, 100))
        self.assertFalse(worker.paper_word_count_within_tolerance(84, 100))
        self.assertFalse(worker.paper_word_count_within_tolerance(116, 100))

    def test_locked_paper_trim_prefers_sentences_and_preserves_latin_tokens(self) -> None:
        sentence_trimmed = worker.trim_locked_paper_content_to_word_limit(
            "First complete sentence. Second complete sentence. Tail token",
            4,
            min_word_count=3,
        )
        self.assertEqual("First complete sentence.", sentence_trimmed)
        self.assertEqual(3, worker.count_text_words(sentence_trimmed))

        token_trimmed = worker.trim_locked_paper_content_to_word_limit(
            "alpha beta supercalifragilistic gamma",
            3,
            min_word_count=3,
        )
        self.assertEqual("alpha beta supercalifragilistic", token_trimmed)
        self.assertNotIn(" supercalifragilis ", f" {token_trimmed} ")

    def test_locked_paper_trim_uses_complete_boundaries_even_below_minimum(self) -> None:
        sentence_trimmed = worker.trim_locked_paper_content_to_word_limit(
            "第一句完整。第二句也完整。第三部分继续考察，后续仍在分析相关影响因素",
            19,
            min_word_count=14,
            fallback_title="研究背景",
        )
        self.assertEqual("第一句完整。第二句也完整。", sentence_trimmed)
        self.assertEqual(11, worker.count_text_words(sentence_trimmed))

        clause_trimmed = worker.trim_locked_paper_content_to_word_limit(
            "研究聚焦用户需求变化，后续内容仍在继续分析平台机制与传播效果",
            12,
            min_word_count=8,
            fallback_title="用户需求",
        )
        self.assertEqual("研究聚焦用户需求变化。", clause_trimmed)
        self.assertEqual(10, worker.count_text_words(clause_trimmed))

        within_cap = worker.trim_locked_paper_content_to_word_limit(
            "研究结论仍需结合样本检验；",
            50,
            fallback_title="研究结论",
        )
        self.assertEqual("研究结论仍需结合样本检验。", within_cap)

        within_group = {"chapter-1": "研究结论仍需结合样本检验，"}
        worker.hard_cap_locked_outline_group_words(
            [{"id": "chapter-1", "title": "研究结论", "level": 1, "target_words": 20}],
            within_group,
        )
        self.assertEqual("研究结论仍需结合样本检验。", within_group["chapter-1"])

    def test_locked_paper_trim_replaces_unbroken_chinese_prefix_with_complete_fallback(self) -> None:
        trimmed = worker.trim_locked_paper_content_to_word_limit(
            "本研究持续深入考察短视频平台内容营销影响机制及大学生消费意愿变化" * 4,
            30,
            min_word_count=25,
            fallback_title="研究背景",
        )

        self.assertTrue(trimmed.startswith("本节围绕研究背景展开分析"))
        self.assertTrue(trimmed.endswith("。"))
        self.assertGreaterEqual(worker.count_text_words(trimmed), 25)
        self.assertLessEqual(worker.count_text_words(trimmed), 30)
        self.assertNotRegex(trimmed, r"(?:的|分|考察)[。.!?！？]?$")

    def test_locked_paper_hard_cap_preserves_exact_outline_and_nonempty_content(self) -> None:
        nodes = [
            {"id": "chapter-1", "title": "Introduction", "level": 1, "target_words": 5},
            {"id": "section-1-1", "title": "Background", "level": 2, "target_words": 5},
        ]
        original_chapter = "one two three four five six"
        content_by_id = {
            "chapter-1": original_chapter,
            "section-1-1": "alpha beta gamma delta epsilon zeta eta theta iota kappa",
        }

        actual_words = worker.hard_cap_locked_outline_group_words(nodes, content_by_id)
        _lower_limit, upper_limit = worker.locked_paper_word_count_limits(10)
        self.assertEqual(upper_limit, actual_words)
        self.assertEqual(original_chapter, content_by_id["chapter-1"])
        self.assertEqual(
            "This section examines Background systematically.",
            content_by_id["section-1-1"],
        )

        tree = worker.build_locked_academic_paper_tree(nodes, content_by_id)
        self.assertTrue(worker.locked_academic_paper_outline_matches(tree, nodes))
        self.assertEqual(["chapter-1"], [node["id"] for node in tree])
        self.assertEqual(["section-1-1"], [node["id"] for node in tree[0]["children"]])
        self.assertTrue(all(content_by_id[node["id"]].strip() for node in nodes))

    def test_locked_paper_hard_cap_keeps_chinese_group_within_limits(self) -> None:
        nodes = [
            {"id": "chapter-1", "title": "研究背景", "level": 1, "target_words": 20},
            {"id": "section-1-1", "title": "研究方法", "level": 2, "target_words": 20},
        ]
        content_by_id = {
            "chapter-1": "本研究持续深入考察短视频平台内容营销影响机制及大学生消费意愿变化" * 4,
            "section-1-1": "研究方法用于说明资料来源与分析步骤，",
        }

        actual_words = worker.hard_cap_locked_outline_group_words(nodes, content_by_id)
        lower_limit, upper_limit = worker.locked_paper_word_count_limits(40)
        self.assertGreaterEqual(actual_words, lower_limit)
        self.assertLessEqual(actual_words, upper_limit)
        self.assertTrue(content_by_id["chapter-1"].startswith("本节围绕研究背景展开分析"))
        self.assertTrue(content_by_id["chapter-1"].endswith("。"))
        self.assertEqual("研究方法用于说明资料来源与分析步骤。", content_by_id["section-1-1"])
        self.assertNotRegex(content_by_id["chapter-1"], r"(?:的|分|考察)[。.!?！？]?$")

        tree = worker.build_locked_academic_paper_tree(nodes, content_by_id)
        self.assertTrue(worker.locked_academic_paper_outline_matches(tree, nodes))
        self.assertEqual(["chapter-1"], [node["id"] for node in tree])
        self.assertEqual(["section-1-1"], [node["id"] for node in tree[0]["children"]])
        self.assertTrue(all(content_by_id[node["id"]].strip() for node in nodes))

        boundary_nodes = [
            {"id": "chapter-2", "title": "研究背景", "level": 1, "target_words": 40},
        ]
        boundary_content = {
            "chapter-2": (
                "第一句完整。第二句也完整。第三部分继续考察，"
                + "后续仍在分析相关影响因素" * 10
            ),
        }
        boundary_words = worker.hard_cap_locked_outline_group_words(boundary_nodes, boundary_content)
        boundary_lower, boundary_upper = worker.locked_paper_word_count_limits(40)
        self.assertGreaterEqual(boundary_words, boundary_lower)
        self.assertLessEqual(boundary_words, boundary_upper)
        self.assertTrue(boundary_content["chapter-2"].startswith("第一句完整。第二句也完整。"))
        self.assertIn("本节围绕研究背景展开分析", boundary_content["chapter-2"])
        self.assertNotIn("第三部分继续考察。", boundary_content["chapter-2"])

    def test_locked_outline_spec_rejects_invalid_version_hierarchy_ids_and_count(self) -> None:
        invalid_specs = [
            {"version": 2, "nodes": [{"id": "a", "title": "A", "level": 1}]},
            {"version": 1, "nodes": []},
            {"version": 1, "nodes": [{"id": "a", "title": "A", "level": 2}]},
            {
                "version": 1,
                "nodes": [
                    {"id": "a", "title": "A", "level": 1},
                    {"id": "b", "title": "B", "level": 3},
                ],
            },
            {
                "version": 1,
                "nodes": [
                    {"id": "a", "title": "A", "level": 1},
                    {"id": "a", "title": "B", "level": 1},
                ],
            },
            {
                "version": 1,
                "nodes": [
                    {"id": f"node-{index}", "title": f"Title {index}", "level": 1}
                    for index in range(worker.PAPER_OUTLINE_MAX_NODES + 1)
                ],
            },
        ]
        for spec in invalid_specs:
            with self.subTest(spec=spec):
                with self.assertRaises(worker.WorkerFailure) as raised:
                    worker.parse_academic_paper_outline_spec(spec, 1000)
                self.assertEqual("PAPER_OUTLINE_SPEC_INVALID", raised.exception.code)

    def test_format_settings_map_flat_admin_fields(self) -> None:
        settings = worker.academic_paper_format_settings(
            {
                "page_format_preset": "custom",
                "page_size": "A4",
                "page_orientation": "portrait",
                "page_margin_top_mm": 25.4,
                "page_gutter_mm": 8,
                "title_east_asia_font": "方正小标宋简体",
                "title_size_pt": 22,
                "body_east_asia_font": "宋体",
                "body_size_pt": 12,
                "heading1_east_asia_font": "黑体",
                "heading1_number_format": "第%1章",
                "heading_numbering_enabled": True,
                "toc_enabled": True,
                "cover_enabled": False,
                "header_distance_cm": 1.5,
                "footer_distance_cm": 1.75,
            }
        )
        self.assertEqual("portrait", settings["page"]["orientation"])
        self.assertEqual(25.4, settings["page"]["margins_mm"]["top"])
        self.assertEqual(8, settings["page"]["margins_mm"]["gutter"])
        self.assertEqual("方正小标宋简体", settings["title"]["east_asia_font"])
        self.assertEqual("宋体", settings["body"]["east_asia_font"])
        self.assertEqual("黑体", settings["headings"]["1"]["east_asia_font"])
        self.assertEqual("第%1章", settings["heading_numbering"]["formats"][0])
        self.assertFalse(settings["pagination"]["include_title_page"])
        self.assertEqual(15, settings["header"]["distance_mm"])
        self.assertEqual(17.5, settings["footer"]["distance_mm"])

    def test_disabled_optional_sections_and_cover_are_omitted(self) -> None:
        paper = worker.build_academic_paper_payload(
            {
                "title": "测试论文",
                "abstract": "摘要",
                "keywords": ["关键词"],
                "references": ["参考文献"],
                "acknowledgements": "致谢",
                "appendices": [{"title": "附录", "content": "内容"}],
            },
            [{"title": "绪论", "level": 1, "content": "正文", "children": []}],
            {
                "abstract_enabled": False,
                "keywords_enabled": False,
                "references_enabled": False,
                "acknowledgements_enabled": False,
                "appendix_enabled": False,
                "cover_enabled": False,
                "cover_school": "不应出现的学校",
                "cover_supervisor": "不应出现的导师",
            },
        )
        for key in ("abstract", "keywords", "references", "acknowledgements", "appendices", "institution", "advisor"):
            self.assertNotIn(key, paper)

        covered = worker.build_academic_paper_payload(
            {"title": "测试论文"},
            [],
            {
                "cover_enabled": True,
                "cover_school": "示例大学",
                "cover_supervisor": "李老师",
                "cover_submission_date": "2026年7月",
            },
        )
        self.assertEqual("示例大学", covered["institution"])
        self.assertEqual("李老师", covered["advisor"])
        self.assertEqual("2026年7月", covered["date"])

    async def test_reference_context_extracts_text_and_skips_template(self) -> None:
        reference = worker.WorkerArtifactRef(
            artifact_id=1,
            name="source.md",
            mime_type="text/markdown",
            url="https://assets.test/source.md",
            metadata={"field_name": "reference_materials"},
        )
        template = worker.WorkerArtifactRef(
            artifact_id=2,
            name="template.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            url="https://assets.test/template.docx",
            metadata={"asset_role": "template"},
        )
        payload = self.payload(artifacts=[reference, template])
        with patch.object(
            worker,
            "download_input_artifact",
            new=AsyncMock(return_value="作者：张三\n题名：数字教育研究".encode()),
        ) as download:
            context = await worker.extract_paper_reference_context(payload)
        self.assertEqual(1, download.await_count)
        self.assertIn("source.md", context)
        self.assertIn("数字教育研究", context)
        self.assertNotIn("template.docx", context)

    async def test_references_disabled_does_not_read_or_parse_reference_assets(self) -> None:
        unreadable_reference = worker.WorkerArtifactRef(
            artifact_id=3,
            name="unreadable-reference.pdf",
            mime_type="application/pdf",
            url="https://assets.test/unreadable-reference.pdf",
            metadata={"field_name": "reference_materials"},
        )
        payload = self.payload(
            input_values={
                "topic": "数字教育治理",
                "word_count": 1000,
                "references_enabled": False,
            },
            artifacts=[unreadable_reference],
        )
        model_results = [
            {
                "response": {
                    "text": json.dumps(
                        {
                            "title": "数字教育治理",
                            "abstract": "摘要",
                            "sections": [
                                {"title": "绪论", "target_words": 500},
                                {"title": "结论", "target_words": 500},
                            ],
                        },
                        ensure_ascii=False,
                    )
                }
            },
            {"response": {"text": json.dumps({"content": "绪论正文" * 100}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"content": "结论正文" * 100}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"consistency_notes": ["已检查"]}, ensure_ascii=False)}},
        ]

        with (
            patch.object(worker, "download_input_artifact", new=AsyncMock(side_effect=OSError("unreadable"))) as download,
            patch.object(worker, "extract_paper_reference_text", side_effect=AssertionError("must not parse")) as parse,
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)),
            patch.object(worker, "build_academic_paper_docx", return_value=b"paper-docx"),
            patch.object(worker, "upload_artifact_bytes", new=AsyncMock(return_value={})),
            patch.object(worker, "callback", new=AsyncMock()),
        ):
            await worker.process_academic_paper_run(payload, "worker-no-references", time.perf_counter())

        download.assert_not_awaited()
        parse.assert_not_called()

    def test_docx_reference_parser_reads_paragraphs_and_tables(self) -> None:
        document = Document()
        document.add_paragraph("文献题名：教育数字化研究")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "作者"
        table.cell(0, 1).text = "张三"
        buffer = BytesIO()
        document.save(buffer)

        artifact = worker.WorkerArtifactRef(
            name="reference.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        text = worker.extract_paper_reference_text(artifact, buffer.getvalue())
        self.assertIn("教育数字化研究", text)
        self.assertIn("作者\t张三", text)

    def test_reference_registry_parses_heading_next_line_and_inline_sources(self) -> None:
        context = """说明中的 [1]-[3] 不应被当作文献。\n来源 [1]\n作者甲。题名甲[J]。DOI: 10.1/one\n\n[2] 作者乙。题名乙[M]。\n来源 [3]\n作者丙。题名丙[J]。"""
        registry = worker.parse_paper_reference_registry(context)
        self.assertIsNotNone(registry)
        assert registry is not None
        self.assertEqual([1, 2, 3], registry["ids"])
        self.assertEqual(
            [
                "作者甲。题名甲[J]。DOI: 10.1/one",
                "作者乙。题名乙[M]。",
                "作者丙。题名丙[J]。",
            ],
            registry["bibliography"],
        )
        self.assertEqual("[1] 作者甲。题名甲[J]。DOI: 10.1/one", registry["numbered_bibliography"][0])

    def test_reference_registry_rejects_duplicate_or_non_contiguous_numbers(self) -> None:
        self.assertIsNone(worker.parse_paper_reference_registry("[1] 一\n[1] 重复"))
        self.assertIsNone(worker.parse_paper_reference_registry("来源 [1]\n一\n来源 [3]\n三"))
        self.assertIsNone(worker.parse_paper_reference_registry("作者：张三\n题名：无编号资料"))
        self.assertIsNone(worker.parse_paper_reference_registry("[1] Body starts here\nordinary paragraph"))
        self.assertIsNone(worker.parse_paper_reference_registry("[1] In 2024 this changed.\n[2] In 2023 that changed."))
        self.assertIsNone(worker.parse_paper_reference_registry("来源 [1]\n\n来源 [2]\n作者乙. 题名乙[J]."))

    def test_reference_registry_keeps_multiline_explicit_source(self) -> None:
        registry = worker.parse_paper_reference_registry(
            "Source [1]\nAJZEN I. The theory of planned behavior[J].\n"
            "Organizational Behavior and Human Decision Processes, 1991.\n"
            "DOI: 10.1016/0749-5978(91)90020-T.\n\n"
        )
        assert registry is not None
        self.assertEqual(1, len(registry["bibliography"]))
        self.assertIn("Organizational Behavior", registry["bibliography"][0])
        self.assertIn("10.1016/0749-5978(91)90020-T", registry["bibliography"][0])

    def test_reference_citation_contract_detects_missing_unknown_and_normalizes_markers(self) -> None:
        registry = worker.parse_paper_reference_registry("来源 [1]\n一\n来源 [2]\n二")
        assert registry is not None
        sections = [{"title": "正文", "content": "论断 [[CITE:1]]。未知 [9]。", "children": []}]
        result = worker.validate_academic_paper_citation_contract(sections, registry)
        self.assertFalse(result["reference_contract_valid"])
        self.assertEqual([1, 9], result["citation_ids_used"])
        self.assertEqual([2], result["citation_ids_missing"])
        self.assertEqual([9], result["citation_ids_unknown"])
        repaired = worker.enforce_academic_paper_citation_contract(sections, registry)
        self.assertFalse(repaired["reference_contract_valid"])
        self.assertFalse(repaired["reference_contract_repaired"])
        self.assertEqual([1, 9], repaired["citation_ids_used"])
        self.assertEqual([2], repaired["citation_ids_missing"])
        self.assertNotIn("[[CITE:", sections[0]["content"])
        self.assertIn("[9]", sections[0]["content"])
        self.assertNotIn("[2]", sections[0]["content"])

        malformed = [{"title": "正文", "content": "论断 [[ CITE:x ]]。", "children": []}]
        malformed_result = worker.validate_academic_paper_citation_contract(malformed, registry)
        self.assertFalse(malformed_result["reference_contract_valid"])
        self.assertTrue(malformed_result["citation_markers_malformed"])

    def test_reference_registry_overrides_model_references_and_format_exposes_style(self) -> None:
        registry = worker.parse_paper_reference_registry("来源 [1]\n代码来源")
        assert registry is not None
        plan = worker.parse_academic_paper_plan(
            json.dumps(
                {
                    "title": "题目",
                    "abstract": "摘要",
                    "sections": [{"title": "绪论"}, {"title": "结论"}],
                    "references": ["模型虚构来源"],
                },
                ensure_ascii=False,
            ),
            {"topic": "主题", "references": ["用户覆盖来源"]},
            1000,
            has_reference_material=True,
            reference_registry=registry,
        )
        self.assertEqual(["代码来源"], plan["references"])
        settings = worker.academic_paper_format_settings({"citation_style": "gbt7714_numeric"})
        self.assertEqual("gbt7714_numeric", settings["citation_style"])

    def test_exact_reference_count_is_only_taken_from_explicit_wording(self) -> None:
        self.assertEqual(5, worker.exact_paper_reference_count({"reference_requirements": "文末恰好5条参考文献"}))
        self.assertEqual(12, worker.exact_paper_reference_count({"reference_requirements": "Exactly 12 references"}))
        self.assertIsNone(worker.exact_paper_reference_count({"reference_requirements": "优先近5年的参考文献"}))
        self.assertIsNone(worker.exact_paper_reference_count({"reference_requirements": "正文恰好5个章节"}))

    def test_numeric_citation_contract_only_applies_to_numeric_styles(self) -> None:
        self.assertTrue(worker.academic_paper_uses_numeric_citations({}))
        self.assertTrue(worker.academic_paper_uses_numeric_citations({"citation_style": "gbt7714_numeric"}))
        self.assertTrue(worker.academic_paper_uses_numeric_citations({"citation_style": "IEEE"}))
        self.assertFalse(worker.academic_paper_uses_numeric_citations({"citation_style": "gbt7714_author_year"}))
        self.assertFalse(worker.academic_paper_uses_numeric_citations({"citation_style": "apa7"}))

    def test_citation_style_resolution_aligns_nested_config_and_renderer_settings(self) -> None:
        cases = [
            ({}, "gbt7714_numeric", True),
            ({"format_settings": {"citation_style": "gbt7714_numeric"}}, "gbt7714_numeric", True),
            (
                {"format_settings": json.dumps({"citation_style": "gbt7714_author_year"})},
                "gbt7714_author_year",
                False,
            ),
            (
                {"citation_style": "apa7", "format_settings": {"citation_style": "gbt7714_numeric"}},
                "apa7",
                False,
            ),
        ]
        for values, expected_style, expected_numeric in cases:
            with self.subTest(values=values):
                self.assertEqual(expected_style, worker.academic_paper_citation_style(values))
                self.assertEqual(expected_numeric, worker.academic_paper_uses_numeric_citations(values))
                self.assertEqual(expected_style, worker.academic_paper_format_settings(values)["citation_style"])

    def test_citation_repair_only_allows_marker_changes(self) -> None:
        sections = [
            {
                "id": "section-one",
                "title": "正文",
                "content": "计划行为理论解释消费意向。",
                "children": [],
            }
        ]
        worker.apply_academic_paper_citation_repair(
            json.dumps(
                {"contents": {"section-one": "计划行为理论解释消费意向[[CITE:1]]。"}},
                ensure_ascii=False,
            ),
            sections,
        )
        self.assertIn("[[CITE:1]]", sections[0]["content"])

        with self.assertRaises(worker.WorkerFailure) as raised:
            worker.apply_academic_paper_citation_repair(
                json.dumps(
                    {"contents": {"section-one": "计划行为理论能够完全证明消费行为[[CITE:1]]。"}},
                    ensure_ascii=False,
                ),
                sections,
            )
        self.assertEqual("PAPER_CITATION_REPAIR_CHANGED_CONTENT", raised.exception.code)

        with self.assertRaises(worker.WorkerFailure) as newline_change:
            worker.apply_academic_paper_citation_repair(
                json.dumps(
                    {"contents": {"section-one": "第一段。\n第二段[[CITE:1]]。"}},
                    ensure_ascii=False,
                ),
                [{"id": "section-one", "title": "正文", "content": "第一段。\n\n第二段。", "children": []}],
            )
        self.assertEqual("PAPER_CITATION_REPAIR_CHANGED_CONTENT", newline_change.exception.code)

    async def test_reference_registry_is_enforced_through_process_and_quality_report(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "引用契约测试",
                "word_count": 1000,
                "outline_spec": {
                    "version": 1,
                    "nodes": [{"id": "chapter-1", "title": "正文", "level": 1}],
                },
            }
        )
        model_results = [
            {
                "response": {
                    "text": json.dumps(
                        {"title": "引用契约测试", "abstract": "摘要", "references": ["模型虚构来源"]},
                        ensure_ascii=False,
                    )
                }
            },
            {
                "response": {
                    "text": json.dumps(
                        {"contents": {"chapter-1": "正文" * 500 + " [[CITE:1]]"}},
                        ensure_ascii=False,
                    )
                }
            },
            {"response": {"text": json.dumps({"consistency_notes": ["已检查"]}, ensure_ascii=False)}},
        ]
        callbacks: list[dict[str, object]] = []

        async def record_callback(_: worker.WorkerRunRequest, event_type: str, **kwargs: object) -> None:
            callbacks.append({"event_type": event_type, **kwargs})

        with (
            patch.object(worker, "extract_paper_reference_context", new=AsyncMock(return_value="来源 [1]\n真实来源")),
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "build_academic_paper_docx", return_value=b"reference-docx") as build_docx,
            patch.object(worker, "upload_artifact_bytes", new=AsyncMock(return_value={})),
            patch.object(worker, "callback", new=record_callback),
        ):
            await worker.process_academic_paper_run(payload, "worker-reference", time.perf_counter())

        self.assertIn("[[CITE:n]]", call_model.await_args_list[1].args[2])
        paper = build_docx.call_args.args[0]
        self.assertEqual(["真实来源"], paper["references"])
        self.assertNotIn("[[CITE:", paper["sections"][0]["content"])
        self.assertIn("[1]", paper["sections"][0]["content"])
        self.assertEqual("gbt7714_numeric", build_docx.call_args.args[1]["citation_style"])
        quality = callbacks[-1]["output"]["quality_report"]
        self.assertTrue(quality["citation_contract_enforced"])
        self.assertTrue(quality["reference_contract_valid"])
        self.assertEqual([1], quality["citation_ids_used"])
        self.assertEqual([], quality["citation_ids_missing"])
        self.assertEqual([], quality["citation_ids_unknown"])

    async def test_academic_paper_orchestrates_plan_sections_review_and_docx(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "生成式人工智能对高等教育的影响",
                "paper_title": "生成式人工智能赋能高等教育研究",
                "word_count": 2000,
                "citation_style": "GB/T 7714",
                "body_east_asia_font": "宋体",
                "body_size_pt": 12,
            }
        )
        plan = {
            "title": "模型改写题目",
            "abstract": "初始摘要",
            "keywords": ["生成式人工智能", "高等教育"],
            "sections": [
                {"title": "绪论", "target_words": 1000, "subsections": ["研究背景"]},
                {"title": "结论", "target_words": 1000, "subsections": ["研究结论"]},
            ],
            "references": [],
        }
        section_one = {
            "title": "绪论",
            "content": "研究背景与问题意识。" * 10,
            "children": [{"title": "研究背景", "level": 2, "content": "高等教育正在经历结构变化。" * 5}],
        }
        expanded_section_one = {
            "title": "绪论",
            "content": "扩展后的完整论证内容。" * 100,
            "children": [{"title": "研究背景", "level": 2, "content": "教育结构变化需要机制分析。" * 30}],
        }
        section_two = {
            "title": "结论",
            "content": "本文形成了与研究问题相对应的结论。" * 80,
            "children": [{"title": "研究结论", "level": 2, "content": "技术应用需要教学制度协同。" * 40}],
        }
        review = {
            "final_title": "不应覆盖用户题目",
            "abstract": "本文分析生成式人工智能对高等教育的影响，并提出制度协同路径。",
            "keywords": ["生成式人工智能", "高等教育", "制度协同"],
            "conclusion_adjustments": "",
            "consistency_notes": ["摘要与结论已对齐"],
        }
        model_results = [
            {"response": {"text": json.dumps(plan, ensure_ascii=False)}, "usage": {"total_tokens": 100}},
            {"response": {"text": json.dumps(section_one, ensure_ascii=False)}, "usage": {"total_tokens": 200}},
            {"response": {"text": json.dumps(expanded_section_one, ensure_ascii=False)}, "usage": {"total_tokens": 160}},
            {"response": {"text": json.dumps(section_two, ensure_ascii=False)}, "usage": {"total_tokens": 220}},
            {"response": {"text": json.dumps(review, ensure_ascii=False)}, "usage": {"total_tokens": 80}},
        ]
        callbacks: list[dict[str, object]] = []

        async def record_callback(_: worker.WorkerRunRequest, event_type: str, **kwargs: object) -> None:
            callbacks.append({"event_type": event_type, **kwargs})

        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "build_academic_paper_docx", return_value=b"fake-docx") as build_docx,
            patch.object(
                worker,
                "upload_artifact_bytes",
                new=AsyncMock(
                    return_value={
                        "artifact_id": 777,
                        "url": "https://download.test/paper.docx",
                        "object_key": "internal/secret/key.docx",
                    }
                ),
            ) as upload,
            patch.object(worker, "callback", new=record_callback),
        ):
            await worker.process_academic_paper_run(payload, "worker-paper", time.perf_counter())

        self.assertEqual(5, call_model.await_count)
        self.assertIn("只返回一个 JSON 对象", call_model.await_args_list[0].args[2])
        self.assertIn("第 1/2 个顶层章节", call_model.await_args_list[1].args[2])
        self.assertIn("完整替换版章节", call_model.await_args_list[2].args[2])
        self.assertIn("一致性审校", call_model.await_args_list[4].args[2])
        build_docx.assert_called_once()
        self.assertEqual("宋体", build_docx.call_args.args[1]["body"]["east_asia_font"])
        self.assertEqual(1, upload.await_count)

        completed = callbacks[-1]
        self.assertEqual("succeeded", completed["event_type"])
        output = completed["output"]
        self.assertEqual("生成式人工智能赋能高等教育研究", build_docx.call_args.args[0]["title"])
        self.assertIn("Word 文件已生成", output["result"])
        self.assertEqual(2, output["quality_report"]["section_count"])
        self.assertEqual("academic-paper.docx", output["document"]["name"])
        self.assertNotIn("artifact_id", output["document"])
        self.assertNotIn("url", output["document"])
        self.assertNotIn("object_key", output["document"])

    async def test_locked_outline_ignores_model_structure_retries_missing_and_reports_match(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "数字教育治理",
                "paper_title": "数字教育治理研究",
                "word_count": 1000,
                "outline_spec": {
                    "version": 1,
                    "nodes": [
                        {"id": "chapter-1", "title": "绪论", "level": 1},
                        {"id": "section-1-1", "title": "研究背景", "level": 2},
                        {"id": "chapter-2", "title": "结论", "level": 1},
                    ],
                },
            }
        )
        model_results = [
            {
                "response": {
                    "text": json.dumps(
                        {
                            "title": "模型题目",
                            "abstract": "初始摘要",
                            "keywords": ["数字教育"],
                            "sections": [{"title": "模型伪造目录", "level": 5}],
                            "references": [],
                        },
                        ensure_ascii=False,
                    )
                },
                "usage": {"total_tokens": 100},
            },
            {
                "response": {
                    "text": json.dumps(
                        {
                            "contents": {
                                "chapter-1": "绪论正文。" * 46,
                                "unknown-node": "必须忽略的未知节点。",
                            },
                            "title": "模型试图改标题",
                        },
                        ensure_ascii=False,
                    )
                },
                "usage": {"total_tokens": 200},
            },
            {
                "response": {
                    "text": json.dumps({"contents": {"chapter-2": "结论正文。" * 90}}, ensure_ascii=False)
                },
                "usage": {"total_tokens": 210},
            },
            {
                "response": {
                    "text": json.dumps({"contents": {"section-1-1": "研究背景正文。" * 64}}, ensure_ascii=False)
                },
                "usage": {"total_tokens": 120},
            },
            {
                "response": {
                    "text": json.dumps(
                        {
                            "final_title": "不能覆盖用户题目",
                            "abstract": "最终摘要",
                            "keywords": ["数字教育", "治理"],
                            "conclusion_adjustments": "不能绕过 ID 映射覆盖结论。",
                            "consistency_notes": ["锁定目录已核验"],
                        },
                        ensure_ascii=False,
                    )
                },
                "usage": {"total_tokens": 80},
            },
        ]
        callbacks: list[dict[str, object]] = []

        async def record_callback(_: worker.WorkerRunRequest, event_type: str, **kwargs: object) -> None:
            callbacks.append({"event_type": event_type, **kwargs})

        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "build_academic_paper_docx", return_value=b"locked-docx") as build_docx,
            patch.object(worker, "upload_artifact_bytes", new=AsyncMock(return_value={"artifact_id": 778})),
            patch.object(worker, "callback", new=record_callback),
        ):
            await worker.process_academic_paper_run(payload, "worker-locked", time.perf_counter())

        self.assertEqual(5, call_model.await_count)
        self.assertIn("不要返回 sections", call_model.await_args_list[0].args[2])
        self.assertIn('"contents"', call_model.await_args_list[1].args[2])
        self.assertIn("仅补写这些缺失节点", call_model.await_args_list[3].args[2])
        self.assertIn("不能重写正文或目录", call_model.await_args_list[4].args[2])

        paper = build_docx.call_args.args[0]
        self.assertEqual("数字教育治理研究", paper["title"])
        self.assertEqual(["chapter-1", "chapter-2"], [section["id"] for section in paper["sections"]])
        self.assertEqual(["绪论", "结论"], [section["title"] for section in paper["sections"]])
        self.assertEqual("section-1-1", paper["sections"][0]["children"][0]["id"])
        self.assertEqual("研究背景", paper["sections"][0]["children"][0]["title"])
        self.assertEqual("结论正文。" * 90, paper["sections"][1]["content"])
        self.assertNotIn("unknown-node", json.dumps(paper, ensure_ascii=False))
        self.assertNotIn("模型伪造目录", json.dumps(paper, ensure_ascii=False))

        quality = callbacks[-1]["output"]["quality_report"]
        self.assertTrue(quality["outline_locked"])
        self.assertEqual(3, quality["outline_node_count"])
        self.assertTrue(quality["outline_match"])
        self.assertEqual([], quality["outline_deviations"])

    async def test_locked_outline_corrects_length_with_same_ids(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "治理机制",
                "word_count": 1000,
                "outline_spec": {
                    "version": 1,
                    "nodes": [{"id": "chapter-1", "title": "治理机制", "level": 1}],
                },
            }
        )
        model_results = [
            {"response": {"text": json.dumps({"title": "治理机制", "abstract": "摘要"})}},
            {"response": {"text": json.dumps({"contents": {"chapter-1": "过短"}}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"contents": {"chapter-1": "校正正文" * 225}}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"consistency_notes": ["已检查"]}, ensure_ascii=False)}},
        ]
        callbacks: list[dict[str, object]] = []

        async def record_callback(_: worker.WorkerRunRequest, event_type: str, **kwargs: object) -> None:
            callbacks.append({"event_type": event_type, **kwargs})

        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "build_academic_paper_docx", return_value=b"locked-docx") as build_docx,
            patch.object(worker, "upload_artifact_bytes", new=AsyncMock(return_value={})),
            patch.object(worker, "callback", new=record_callback),
        ):
            await worker.process_academic_paper_run(payload, "worker-correct", time.perf_counter())

        self.assertEqual(4, call_model.await_count)
        self.assertIn("校正篇幅", call_model.await_args_list[2].args[2])
        self.assertIn("不要返回 title、level、children", call_model.await_args_list[2].args[2])
        section = build_docx.call_args.args[0]["sections"][0]
        self.assertEqual("chapter-1", section["id"])
        self.assertEqual("治理机制", section["title"])
        self.assertEqual("校正正文" * 225, section["content"])
        self.assertEqual([], callbacks[-1]["output"]["quality_report"]["outline_deviations"])

    async def test_locked_outline_repeats_correction_until_later_attempt_matches(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "治理机制",
                "word_count": 1000,
                "outline_spec": {
                    "version": 1,
                    "nodes": [{"id": "chapter-1", "title": "治理机制", "level": 1}],
                },
            }
        )
        model_results = [
            {"response": {"text": json.dumps({"title": "治理机制", "abstract": "摘要"})}},
            {"response": {"text": json.dumps({"contents": {"chapter-1": "过短"}}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"contents": {"chapter-1": "仍然过长" * 400}}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"contents": {"chapter-1": "合格正文" * 250}}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"consistency_notes": ["已检查"]}, ensure_ascii=False)}},
        ]
        callbacks: list[dict[str, object]] = []

        async def record_callback(_: worker.WorkerRunRequest, event_type: str, **kwargs: object) -> None:
            callbacks.append({"event_type": event_type, **kwargs})

        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "PAPER_OUTLINE_MAX_CORRECTION_ATTEMPTS", 3),
            patch.object(worker, "build_academic_paper_docx", return_value=b"locked-docx") as build_docx,
            patch.object(worker, "upload_artifact_bytes", new=AsyncMock(return_value={})),
            patch.object(worker, "callback", new=record_callback),
        ):
            await worker.process_academic_paper_run(payload, "worker-repeat-correct", time.perf_counter())

        self.assertEqual(5, call_model.await_count)
        self.assertIn("当前正文（约 1600 字", call_model.await_args_list[3].args[2])
        section = build_docx.call_args.args[0]["sections"][0]
        self.assertEqual("chapter-1", section["id"])
        self.assertEqual("治理机制", section["title"])
        self.assertEqual("合格正文" * 250, section["content"])
        self.assertEqual([], callbacks[-1]["output"]["quality_report"]["outline_deviations"])

    async def test_locked_outline_hard_caps_oversize_after_correction_attempts_exhausted(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "治理机制",
                "word_count": 1000,
                "outline_spec": {
                    "version": 1,
                    "nodes": [{"id": "chapter-1", "title": "治理机制", "level": 1}],
                },
            }
        )
        oversized = json.dumps({"contents": {"chapter-1": "仍然过长" * 400}}, ensure_ascii=False)
        model_results = [
            {"response": {"text": json.dumps({"title": "治理机制", "abstract": "摘要"})}},
            {"response": {"text": oversized}},
            {"response": {"text": oversized}},
            {"response": {"text": oversized}},
            {"response": {"text": oversized}},
            {"response": {"text": json.dumps({"consistency_notes": ["已检查"]}, ensure_ascii=False)}},
        ]
        callbacks: list[dict[str, object]] = []

        async def record_callback(_: worker.WorkerRunRequest, event_type: str, **kwargs: object) -> None:
            callbacks.append({"event_type": event_type, **kwargs})

        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "PAPER_OUTLINE_MAX_CORRECTION_ATTEMPTS", 3),
            patch.object(worker, "build_academic_paper_docx", return_value=b"locked-docx") as build_docx,
            patch.object(worker, "upload_artifact_bytes", new=AsyncMock(return_value={})),
            patch.object(worker, "callback", new=record_callback),
        ):
            await worker.process_academic_paper_run(payload, "worker-exhaust-correct", time.perf_counter())

        self.assertEqual(6, call_model.await_count)
        quality = callbacks[-1]["output"]["quality_report"]
        self.assertTrue(quality["outline_match"])
        self.assertEqual([], quality["outline_deviations"])
        self.assertLessEqual(quality["deviation_percent"], 15)
        section = build_docx.call_args.args[0]["sections"][0]
        _lower_limit, upper_limit = worker.locked_paper_word_count_limits(section["target_words"])
        self.assertLessEqual(worker.count_text_words(section["content"]), upper_limit)
        self.assertTrue(section["content"].strip())
        self.assertEqual(("chapter-1", 1), (section["id"], section["level"]))

    async def test_locked_outline_reports_persistent_undersize_after_corrections(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "Persistent undersize",
                "word_count": 1000,
                "outline_spec": {
                    "version": 1,
                    "nodes": [{"id": "chapter-1", "title": "Only chapter", "level": 1}],
                },
            }
        )
        undersized = json.dumps(
            {"contents": {"chapter-1": " ".join(f"word-{index}" for index in range(100))}}
        )
        model_results = [
            {"response": {"text": json.dumps({"title": "Persistent undersize", "abstract": "Summary"})}},
            {"response": {"text": undersized}},
            {"response": {"text": undersized}},
            {"response": {"text": undersized}},
            {"response": {"text": undersized}},
            {"response": {"text": json.dumps({"consistency_notes": ["Reviewed"]})}},
        ]
        callbacks: list[dict[str, object]] = []

        async def record_callback(_: worker.WorkerRunRequest, event_type: str, **kwargs: object) -> None:
            callbacks.append({"event_type": event_type, **kwargs})

        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "PAPER_OUTLINE_MAX_CORRECTION_ATTEMPTS", 3),
            patch.object(worker, "build_academic_paper_docx", return_value=b"locked-docx"),
            patch.object(worker, "upload_artifact_bytes", new=AsyncMock(return_value={})),
            patch.object(worker, "callback", new=record_callback),
        ):
            await worker.process_academic_paper_run(payload, "worker-undersize-correct", time.perf_counter())

        self.assertEqual(6, call_model.await_count)
        quality = callbacks[-1]["output"]["quality_report"]
        self.assertTrue(quality["outline_match"])
        self.assertEqual(1, len(quality["outline_deviations"]))
        deviation = quality["outline_deviations"][0]
        self.assertEqual("chapter-1", deviation["node_id"])
        self.assertEqual(100, deviation["actual_word_count"])
        self.assertLess(deviation["deviation_percent"], 0)

    async def test_locked_outline_keeps_planned_abstract_when_review_would_break_overall_tolerance(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "治理机制",
                "word_count": 1000,
                "outline_spec": {
                    "version": 1,
                    "nodes": [{"id": "chapter-1", "title": "治理机制", "level": 1}],
                },
            }
        )
        model_results = [
            {"response": {"text": json.dumps({"title": "治理机制", "abstract": "摘要"}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"contents": {"chapter-1": "正文" * 500}}, ensure_ascii=False)}},
            {
                "response": {
                    "text": json.dumps(
                        {"abstract": "过长摘要" * 100, "consistency_notes": ["已检查"]},
                        ensure_ascii=False,
                    )
                }
            },
        ]
        callbacks: list[dict[str, object]] = []

        async def record_callback(_: worker.WorkerRunRequest, event_type: str, **kwargs: object) -> None:
            callbacks.append({"event_type": event_type, **kwargs})

        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "build_academic_paper_docx", return_value=b"locked-docx") as build_docx,
            patch.object(worker, "upload_artifact_bytes", new=AsyncMock(return_value={})),
            patch.object(worker, "callback", new=record_callback),
        ):
            await worker.process_academic_paper_run(payload, "worker-abstract-guard", time.perf_counter())

        self.assertEqual(3, call_model.await_count)
        self.assertEqual("摘要", build_docx.call_args.args[0]["abstract"])
        quality = callbacks[-1]["output"]["quality_report"]
        self.assertLessEqual(abs(quality["deviation_percent"]), 15)
        self.assertIn("终审摘要会使总字数超出容差", " ".join(quality["consistency_notes"]))

    async def test_locked_outline_fails_when_missing_nodes_remain_after_one_retry(self) -> None:
        payload = self.payload(
            input_values={
                "topic": "数字治理",
                "word_count": 1000,
                "outline_spec": {
                    "version": 1,
                    "nodes": [{"id": "chapter-1", "title": "绪论", "level": 1}],
                },
            }
        )
        model_results = [
            {"response": {"text": json.dumps({"title": "数字治理", "abstract": "摘要"})}},
            {"response": {"text": json.dumps({"contents": {"unknown": "未知节点"}}, ensure_ascii=False)}},
            {"response": {"text": json.dumps({"contents": {}}, ensure_ascii=False)}},
        ]
        with (
            patch.object(worker, "call_model_proxy", new=AsyncMock(side_effect=model_results)) as call_model,
            patch.object(worker, "callback", new=AsyncMock()),
        ):
            with self.assertRaises(worker.WorkerFailure) as raised:
                await worker.process_academic_paper_run(payload, "worker-missing", time.perf_counter())

        self.assertEqual("PAPER_OUTLINE_CONTENT_MISSING", raised.exception.code)
        self.assertEqual(3, call_model.await_count)


if __name__ == "__main__":
    unittest.main()
