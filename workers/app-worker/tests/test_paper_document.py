from __future__ import annotations

import io
import sys
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from sub2api_worker.paper_document import (  # noqa: E402
    PAGE_SIZES_MM,
    build_academic_paper_docx,
    count_academic_paper_words,
    summarize_academic_paper_format,
)
from sub2api_worker.main import extract_docx_reference_text  # noqa: E402


class AcademicPaperDocumentTests(unittest.TestCase):
    @staticmethod
    def direct_page_break_count(paragraph) -> int:
        return len(paragraph._p.xpath("./w:pPr/w:pageBreakBefore"))

    @staticmethod
    def paragraph_with_text(doc, text: str, style_name: str | None = None):
        return next(
            paragraph
            for paragraph in doc.paragraphs
            if paragraph.text == text and (style_name is None or paragraph.style.name == style_name)
        )

    def sample_paper(self) -> dict[str, object]:
        return {
            "title": "人工智能辅助城市治理研究",
            "subtitle": "基于多源数据的分析",
            "author": "张三",
            "institution": "示例大学",
            "department": "计算机学院",
            "major": "计算机科学与技术",
            "student_id": "20260001",
            "advisor": "李老师",
            "date": "2026年7月",
            "abstract": "本文研究**人工智能**在城市治理中的应用。\n\n研究结果表明该方法具有可行性。",
            "keywords": ["人工智能", "城市治理", "多源数据"],
            "sections": [
                {
                    "title": "# 1 绪论",
                    "content": "```text\n第一段介绍研究背景。\n```\n\n第二段包含 **重点内容**。",
                    "children": [
                        {
                            "title": "1.1 研究背景",
                            "content": "城市治理正在形成 data-driven decision making 模式。",
                            "children": [
                                {
                                    "title": "1.1.1 技术基础",
                                    "content": "AI systems can process multi-source data.",
                                }
                            ],
                        }
                    ],
                },
                {"title": "2 结论", "content": "研究目标已经完成。"},
            ],
            "references": [
                "张三. 人工智能研究[M]. 北京: 示例出版社, 2026.",
                {"citation": "Smith J. Urban AI[J]. Example Journal, 2025, 1(2): 1-10."},
            ],
            "acknowledgements": "感谢指导教师和参与调研的工作人员。",
            "appendices": [{"title": "附录 A 调研提纲", "content": "调研问题与数据说明。"}],
        }

    def test_builds_valid_docx_with_styles_numbering_toc_and_fields(self) -> None:
        settings = {
            "page_size": "A4",
            "page_margin_left_cm": 3.2,
            "body_east_asia_font": "仿宋",
            "body_latin_font": "Arial",
            "body_size_pt": 11,
            "body_line_spacing_mode": "exact",
            "body_line_spacing_value": 22,
            "heading1_east_asia_font": "方正小标宋简体",
            "heading1_size_pt": 18,
            "heading1_alignment": "center",
            "heading5_size_pt": 10.5,
            "toc_levels": 5,
            "header_enabled": True,
            "header_text": "示例大学本科论文",
            "page_number_enabled": True,
            "page_number_position": "footer",
            "page_number_start": 3,
            "page_number_format": "lowerRoman",
        }
        data = build_academic_paper_docx(self.sample_paper(), settings)
        self.assertTrue(data.startswith(b"PK"))
        self.assertGreater(len(data), 10_000)

        doc = Document(io.BytesIO(data))
        all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("人工智能辅助城市治理研究", all_text)
        self.assertIn("参考文献", all_text)
        self.assertIn("致谢", all_text)
        self.assertIn("附录 A 调研提纲", all_text)
        self.assertNotIn("```", all_text)
        self.assertNotIn("**", all_text)
        self.assertNotIn("# 1 绪论", all_text)

        heading_one = next(
            paragraph for paragraph in doc.paragraphs if paragraph.style.name == "Heading 1" and paragraph.text == "绪论"
        )
        self.assertIsNotNone(heading_one._p.pPr.numPr)
        self.assertEqual(
            "方正小标宋简体",
            doc.styles["Heading 1"]._element.rPr.rFonts.get(qn("w:eastAsia")),
        )
        self.assertAlmostEqual(18.0, doc.styles["Heading 1"].font.size.pt, places=1)
        self.assertEqual("仿宋", doc.styles["Normal"]._element.rPr.rFonts.get(qn("w:eastAsia")))
        self.assertEqual("Arial", doc.styles["Normal"]._element.rPr.rFonts.get(qn("w:ascii")))

        self.assertAlmostEqual(32.0, doc.sections[0].left_margin.mm, places=1)
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            settings_xml = archive.read("word/settings.xml").decode("utf-8")
            numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
            footer_xml = archive.read("word/footer1.xml").decode("utf-8")
            self.assertIn('TOC \\o "1-5" \\h \\z \\u', document_xml)
            self.assertIn("研究背景", document_xml)
            self.assertIn('w:updateFields w:val="true"', settings_xml)
            self.assertIn("AcademicHeadingNumbering", numbering_xml)
            self.assertIn("第%1章", numbering_xml)
            self.assertIn("> PAGE <", footer_xml)
            self.assertIn('w:pgNumType w:start="3" w:fmt="lowerRoman"', document_xml)

    def test_numeric_reference_style_adds_continuous_numbers_without_changing_citations(self) -> None:
        references = [
            "[9] AJZEN I. The theory of planned behavior[J]. DOI: 10.1016/0749-5978(91)90020-T.",
            {"citation": "MEHRABIAN A, RUSSELL J A. An Approach to Environmental Psychology[M]."},
            "【42】 HOLLEBEEK L D, MACKY K. Digital content marketing[J]. DOI: 10.1016/j.intmar.2018.07.003.",
            "LOU C, YUAN S. Influencer marketing[J]. DOI: 10.1080/15252019.2018.1533501.",
            "SOKOLOVA K, KEFI H. Instagram and YouTube bloggers promote it[J]. DOI: 10.1016/j.jretconser.2019.01.011.",
        ]
        data = build_academic_paper_docx(
            {
                "title": "Reference numbering test",
                "sections": [{"title": "Body", "content": "Body text [1]."}],
                "references": references,
            },
            {
                "citation_style": "gbt7714_numeric",
                "toc": {"enabled": False},
                "pagination": {"include_title_page": False},
                "page_number": {"enabled": False},
            },
        )
        doc = Document(io.BytesIO(data))
        rendered = [paragraph.text for paragraph in doc.paragraphs if paragraph.style.name == "Academic Reference"]

        self.assertEqual(5, len(rendered))
        self.assertEqual([f"[{index}]" for index in range(1, 6)], [text.split(" ", 1)[0] for text in rendered])
        self.assertNotIn("[9]", rendered[0])
        self.assertNotIn("【42】", rendered[2])
        self.assertIn("10.1016/0749-5978(91)90020-T.", rendered[0])
        self.assertIn("10.1016/j.intmar.2018.07.003.", rendered[2])
        self.assertIn("10.1016/j.jretconser.2019.01.011.", rendered[4])

    def test_author_year_reference_style_preserves_original_text(self) -> None:
        citation = "AJZEN I. The theory of planned behavior[J]."
        data = build_academic_paper_docx(
            {
                "title": "Reference style test",
                "sections": [{"title": "Body", "content": "Body text."}],
                "references": [citation],
            },
            {
                "citation_style": "gbt7714_author_year",
                "toc": {"enabled": False},
                "pagination": {"include_title_page": False},
                "page_number": {"enabled": False},
            },
        )
        doc = Document(io.BytesIO(data))
        rendered = [paragraph.text for paragraph in doc.paragraphs if paragraph.style.name == "Academic Reference"]
        self.assertEqual([citation], rendered)

    def test_extract_docx_reference_text_preserves_interleaved_xml_block_order(self) -> None:
        source = Document()
        source.add_paragraph("Paragraph before first table")
        first_table = source.add_table(rows=1, cols=2)
        first_table.cell(0, 0).text = "First table author"
        first_table.cell(0, 1).text = "First table DOI"
        source.add_paragraph("Paragraph between tables")
        second_table = source.add_table(rows=1, cols=2)
        second_table.cell(0, 0).text = "Second table author"
        second_table.cell(0, 1).text = "Second table DOI"
        source.add_paragraph("Paragraph after second table")
        stream = io.BytesIO()
        source.save(stream)

        extracted = extract_docx_reference_text(stream.getvalue())

        self.assertEqual(
            "\n".join(
                [
                    "Paragraph before first table",
                    "First table author\tFirst table DOI",
                    "Paragraph between tables",
                    "Second table author\tSecond table DOI",
                    "Paragraph after second table",
                ]
            ),
            extracted,
        )

    def test_front_matter_page_boundaries_are_idempotent(self) -> None:
        paper = {
            "title": "Boundary Test",
            "abstract": "Abstract body.",
            "sections": [{"title": "First chapter", "content": "Chapter body."}],
        }

        for toc_before_abstract in (False, True):
            with self.subTest(toc_before_abstract=toc_before_abstract):
                data = build_academic_paper_docx(
                    paper,
                    {
                        "toc": {"before_abstract": toc_before_abstract, "title": "Contents"},
                        "labels": {"abstract": "Abstract"},
                        "page_number": {"enabled": False},
                    },
                )
                doc = Document(io.BytesIO(data))
                front_matter = [
                    self.paragraph_with_text(doc, "Abstract", "Heading 1"),
                    self.paragraph_with_text(doc, "Contents", "Academic TOC Title"),
                    self.paragraph_with_text(doc, "First chapter", "Heading 1"),
                ]

                for paragraph in front_matter:
                    self.assertEqual(1, self.direct_page_break_count(paragraph))
                self.assertFalse(doc._element.body.xpath('.//w:br[@w:type="page"]'))

    def test_front_matter_page_break_switches_can_all_be_disabled(self) -> None:
        data = build_academic_paper_docx(
            {
                "title": "No Boundary Test",
                "abstract": "Abstract body.",
                "sections": [{"title": "First chapter", "content": "Chapter body."}],
            },
            {
                "toc": {
                    "title": "Contents",
                    "page_break_before": False,
                    "page_break_after": False,
                },
                "labels": {"abstract": "Abstract"},
                "pagination": {
                    "include_title_page": True,
                    "title_page_break_after": False,
                    "abstract_page_break_after": False,
                    "toc_page_break_after": False,
                    "chapter_page_break_before": False,
                },
                "page_number": {"enabled": False},
            },
        )
        doc = Document(io.BytesIO(data))

        paragraphs = [
            self.paragraph_with_text(doc, "Abstract", "Heading 1"),
            self.paragraph_with_text(doc, "Contents", "Academic TOC Title"),
            self.paragraph_with_text(doc, "First chapter", "Heading 1"),
        ]
        for paragraph in paragraphs:
            self.assertEqual(0, self.direct_page_break_count(paragraph))
        self.assertFalse(doc._element.body.xpath('.//w:br[@w:type="page"]'))

    def test_supports_named_and_custom_page_sizes(self) -> None:
        paper = {"title": "Page Geometry Test", "sections": [{"title": "Body", "content": "Text."}]}
        for page_size in ("A3", "A4", "A5", "B5", "B5_JIS", "LETTER", "LEGAL"):
            with self.subTest(page_size=page_size):
                data = build_academic_paper_docx(
                    paper,
                    {
                        "page_size": page_size,
                        "toc_enabled": False,
                        "cover_enabled": False,
                        "page_number_enabled": False,
                    },
                )
                doc = Document(io.BytesIO(data))
                expected_width, expected_height = PAGE_SIZES_MM[page_size]
                self.assertAlmostEqual(expected_width, doc.sections[0].page_width.mm, delta=0.1)
                self.assertAlmostEqual(expected_height, doc.sections[0].page_height.mm, delta=0.1)

        custom = build_academic_paper_docx(
            paper,
            {
                "page_size": "CUSTOM",
                "page_width_mm": 180,
                "page_height_mm": 260,
                "orientation": "landscape",
                "page_margin_top_cm": 1.8,
                "page_margin_bottom_cm": 1.9,
                "page_margin_left_cm": 2.0,
                "page_margin_right_cm": 2.1,
                "toc_enabled": False,
                "cover_enabled": False,
            },
        )
        custom_doc = Document(io.BytesIO(custom))
        section = custom_doc.sections[0]
        self.assertAlmostEqual(260.0, section.page_width.mm, delta=0.1)
        self.assertAlmostEqual(180.0, section.page_height.mm, delta=0.1)
        self.assertAlmostEqual(18.0, section.top_margin.mm, delta=0.1)
        self.assertAlmostEqual(19.0, section.bottom_margin.mm, delta=0.1)
        self.assertAlmostEqual(20.0, section.left_margin.mm, delta=0.1)
        self.assertAlmostEqual(21.0, section.right_margin.mm, delta=0.1)

    def test_counts_text_summarizes_format_and_uses_real_template(self) -> None:
        paper = self.sample_paper()
        counts = count_academic_paper_words(paper)
        self.assertGreater(counts["word_count"], 50)
        self.assertGreater(counts["cjk_characters"], 40)
        self.assertGreater(counts["latin_words"], 10)
        self.assertEqual(4, counts["section_count"])

        summary = summarize_academic_paper_format(
            {
                "page_size": "LEGAL",
                "orientation": "landscape",
                "body_line_spacing_mode": "multiple",
                "body_line_spacing_value": 2,
                "heading5_size_pt": 9,
            }
        )
        self.assertEqual("LEGAL", summary["page"]["size"])
        self.assertEqual("landscape", summary["page"]["orientation"])
        self.assertEqual("multiple:2", summary["body"]["line_spacing"])
        self.assertEqual(9.0, summary["headings"]["5"]["size_pt"])

        apa = summarize_academic_paper_format({"format_preset": "apa_english"})
        self.assertEqual("LETTER", apa["page"]["size"])
        self.assertEqual("Times New Roman", apa["body"]["latin_font"])
        self.assertEqual("multiple:2", apa["body"]["line_spacing"])
        self.assertEqual(1.27, apa["body"]["first_line_indent_cm"])
        journal = summarize_academic_paper_format({"format_preset": "cn_journal"})
        self.assertEqual(10.5, journal["body"]["size_pt"])
        self.assertFalse(journal["toc"]["enabled"])
        graduate = summarize_academic_paper_format({"format_preset": "graduate_thesis"})
        self.assertEqual("exact:20pt", graduate["body"]["line_spacing"])
        global_fonts = summarize_academic_paper_format(
            {"default_east_asia_font": "楷体", "default_latin_font": "Arial"}
        )
        self.assertEqual("楷体", global_fonts["body"]["east_asia_font"])
        self.assertEqual("Arial", global_fonts["headings"]["3"]["latin_font"])

        template = Document()
        template.styles.add_style("University Template Marker", 1)
        template.styles["Normal"].font.name = "Courier New"
        template.styles["Normal"].font.size = Pt(13)
        template.styles["Heading 1"].font.name = "Georgia"
        template.styles["Heading 1"].font.size = Pt(17)
        template.sections[0].page_width = Mm(148)
        template.sections[0].page_height = Mm(210)
        template.sections[0].header.paragraphs[0].text = "TEMPLATE HEADER"
        template.add_paragraph("THIS TEMPLATE BODY MUST BE REMOVED")
        template_stream = io.BytesIO()
        template.save(template_stream)
        data = build_academic_paper_docx(
            {"title": "Template Test", "sections": [{"title": "Content", "content": "Body."}]},
            {
                "format_preset": "uploaded_template",
                "toc_enabled": False,
                "cover_enabled": False,
                "body_size_pt": 11,
                "page_orientation": "landscape",
            },
            template_bytes=template_stream.getvalue(),
        )
        result = Document(io.BytesIO(data))
        self.assertIn("University Template Marker", [style.name for style in result.styles])
        self.assertNotIn("THIS TEMPLATE BODY MUST BE REMOVED", "\n".join(p.text for p in result.paragraphs))
        self.assertEqual("Courier New", result.styles["Normal"].font.name)
        self.assertAlmostEqual(11.0, result.styles["Normal"].font.size.pt, places=1)
        self.assertEqual("Georgia", result.styles["Heading 1"].font.name)
        self.assertAlmostEqual(17.0, result.styles["Heading 1"].font.size.pt, places=1)
        self.assertAlmostEqual(210.0, result.sections[0].page_width.mm, delta=0.1)
        self.assertAlmostEqual(148.0, result.sections[0].page_height.mm, delta=0.1)
        self.assertEqual("TEMPLATE HEADER", result.sections[0].header.paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
