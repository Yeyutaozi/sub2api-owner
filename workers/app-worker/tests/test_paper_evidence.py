from __future__ import annotations

import io
import sys
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document


sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from sub2api_worker.paper_evidence import (  # noqa: E402
    CitationOccurrence,
    EvidenceAssertion,
    EvidenceSource,
    EvidenceSourceIdentityError,
    UnsupportedEvidenceFormat,
    build_evidence_corpus,
    extract_docx_evidence_blocks,
    extract_pdf_evidence_blocks,
    extract_text_evidence_blocks,
    infer_reference_id_from_artifact_name,
    normalize_evidence_quote,
    normalize_evidence_text,
    validate_evidence_source_identities,
    validate_evidence_quote,
    verify_evidence_audit,
)


class PaperEvidenceTests(unittest.TestCase):
    def test_normalizes_unicode_whitespace_and_outer_quote_delimiters(self) -> None:
        self.assertEqual("ABC 研究结果", normalize_evidence_text("ＡＢＣ\u200b\n  研究\u00ad结果"))
        self.assertEqual("证据原句", normalize_evidence_quote("“ 证据原句 ”"))

    def test_extracts_nonempty_text_lines_with_original_line_numbers(self) -> None:
        blocks = extract_text_evidence_blocks(
            "第一行\n\n第三行包含充分的证据内容。\n".encode(),
            artifact_name="[2] source.txt",
            reference_id=2,
            source_key=8,
        )

        self.assertEqual(["第一行", "第三行包含充分的证据内容。"], [block.text for block in blocks])
        self.assertEqual([1, 3], [block.locator.index for block in blocks])
        self.assertTrue(all(block.locator.kind == "line" for block in blocks))
        self.assertTrue(all(block.reference_id == 2 for block in blocks))
        self.assertEqual(blocks[0].chunk_id, extract_text_evidence_blocks(
            "第一行\n\n第三行包含充分的证据内容。\n".encode(),
            artifact_name="[2] source.txt",
            reference_id=2,
            source_key=8,
        )[0].chunk_id)

    def test_extracts_docx_paragraphs_and_table_rows_in_document_order(self) -> None:
        document = Document()
        document.add_paragraph("段落一包含可核验内容。")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "表格证据甲"
        table.cell(0, 1).text = "表格证据乙"
        document.add_paragraph("段落三包含另一条证据。")
        output = io.BytesIO()
        document.save(output)

        blocks = extract_docx_evidence_blocks(
            output.getvalue(),
            artifact_name="paper.docx",
            reference_id=1,
        )

        self.assertEqual(
            ["段落一包含可核验内容。", "表格证据甲\t表格证据乙", "段落三包含另一条证据。"],
            [block.text for block in blocks],
        )
        self.assertEqual([1, 2, 3], [block.locator.index for block in blocks])
        self.assertTrue(all(block.locator.kind == "paragraph" for block in blocks))

    def test_extracts_pdf_pages_and_preserves_page_numbers(self) -> None:
        pages = [
            SimpleNamespace(extract_text=lambda: "第一页有足够长的原文证据。"),
            SimpleNamespace(extract_text=lambda: ""),
            SimpleNamespace(extract_text=lambda: "第三页的研究结论可以核验。"),
        ]
        with patch("pypdf.PdfReader", return_value=SimpleNamespace(pages=pages)):
            blocks = extract_pdf_evidence_blocks(
                b"fake-pdf",
                artifact_name="source.pdf",
                reference_id=3,
                source_key="artifact-3",
            )

        self.assertEqual([1, 3], [block.locator.index for block in blocks])
        self.assertTrue(all(block.locator.kind == "page" for block in blocks))

    def test_infers_reference_id_from_supported_file_names(self) -> None:
        self.assertEqual(1, infer_reference_id_from_artifact_name("[1] paper.pdf"))
        self.assertEqual(2, infer_reference_id_from_artifact_name("2-paper.docx"))
        self.assertEqual(3, infer_reference_id_from_artifact_name("ref-3.txt"))
        self.assertIsNone(infer_reference_id_from_artifact_name("paper-2026.txt"))

    def test_explicit_reference_id_takes_priority_over_name_and_sections(self) -> None:
        corpus = build_evidence_corpus(
            [
                EvidenceSource(
                    artifact_name="[2] source.txt",
                    reference_id=7,
                    data="Source [3]\n这条原文仍应属于显式编号七。".encode(),
                )
            ]
        )

        self.assertTrue(corpus.blocks)
        self.assertTrue(all(block.reference_id == 7 for block in corpus.blocks))

    def test_assigns_reference_sections_and_removes_marker_lines(self) -> None:
        corpus = build_evidence_corpus(
            [
                EvidenceSource(
                    artifact_name="references.txt",
                    artifact_id=91,
                    data=(
                        "资料说明\n"
                        "Source [1]\n"
                        "第一篇文献提供了这一条可核验证据。\n"
                        "来源【2】：第二篇文献首条证据内容充分。\n"
                        "第二篇文献的补充证据也保留。"
                    ).encode(),
                )
            ]
        )

        self.assertEqual([None, 1, 2, 2], [block.reference_id for block in corpus.blocks])
        self.assertNotIn("Source [1]", [block.text for block in corpus.blocks])
        self.assertEqual((1, 2), corpus.reference_ids)

    def test_validates_source_identity_by_doi_and_title_fallback(self) -> None:
        doi_corpus = build_evidence_corpus(
            [
                EvidenceSource(
                    artifact_name="[1] governance.txt",
                    artifact_id=11,
                    data=(
                        "数字治理研究\nDOI: 10.1234/governance.2026\n"
                        "本文讨论数字治理对公共服务的影响。"
                    ).encode(),
                )
            ]
        )
        doi_records = validate_evidence_source_identities(
            [
                {
                    "id": 1,
                    "citation": "作者甲. 数字治理研究[J]. 治理学报, 2026. DOI: 10.1234/governance.2026.",
                }
            ],
            doi_corpus,
        )
        self.assertEqual("doi", doi_records[0]["match_method"])
        self.assertEqual("10.1234/governance.2026", doi_records[0]["matched_doi"])

        title_corpus = build_evidence_corpus(
            [
                EvidenceSource(
                    artifact_name="[1] collaboration.txt",
                    artifact_id=12,
                    data="协同治理机制研究\n作者乙\n本文分析资源配置机制。".encode(),
                )
            ]
        )
        title_records = validate_evidence_source_identities(
            [{"id": 1, "citation": "作者乙. 协同治理机制研究[J]. 公共管理学报, 2025."}],
            title_corpus,
        )
        self.assertEqual("title", title_records[0]["match_method"])
        self.assertEqual("协同治理机制研究", title_records[0]["matched_title"])

    def test_rejects_conflicting_or_unverifiable_source_identity(self) -> None:
        conflicting = build_evidence_corpus(
            [
                EvidenceSource(
                    artifact_name="[1] wrong.txt",
                    data="另一篇研究\nDOI: 10.9999/wrong.2025\n作者甲".encode(),
                )
            ]
        )
        with self.assertRaisesRegex(EvidenceSourceIdentityError, "DOI conflicts"):
            validate_evidence_source_identities(
                [{"id": 1, "citation": "作者甲. 数字治理研究[J]. DOI: 10.1234/right.2026."}],
                conflicting,
            )

        same_author_wrong_title = build_evidence_corpus(
            [
                EvidenceSource(
                    artifact_name="[1] wrong-title.txt",
                    data="作者甲\n完全不同的论文标题\n正文内容。".encode(),
                )
            ]
        )
        with self.assertRaisesRegex(EvidenceSourceIdentityError, "cannot be tied"):
            validate_evidence_source_identities(
                [{"id": 1, "citation": "作者甲. 数字治理研究[J]. 治理学报, 2026."}],
                same_author_wrong_title,
            )

    def test_rejects_numbered_evidence_outside_bibliography(self) -> None:
        corpus = build_evidence_corpus(
            [
                EvidenceSource(artifact_name="[1] first.txt", data="第一篇论文标题和正文。".encode()),
                EvidenceSource(artifact_name="[2] extra.txt", data="额外论文标题和正文。".encode()),
            ]
        )
        with self.assertRaisesRegex(EvidenceSourceIdentityError, "not present in bibliography"):
            validate_evidence_source_identities(
                [{"id": 1, "citation": "作者甲. 第一篇论文标题[J]. 治理学报, 2026."}],
                corpus,
            )

        unnumbered = build_evidence_corpus(
            [
                EvidenceSource(
                    artifact_name="[1] first.txt",
                    data="第一篇论文标题\n第一篇正文。".encode(),
                ),
                EvidenceSource(
                    artifact_name="unmapped.txt",
                    data="未编号的独立资料不应进入严格证据上下文。".encode(),
                ),
            ]
        )
        with self.assertRaisesRegex(EvidenceSourceIdentityError, "not assigned"):
            validate_evidence_source_identities(
                [{"id": 1, "citation": "作者甲. 第一篇论文标题[J]. 治理学报, 2026."}],
                unnumbered,
            )

    def test_rejects_unsupported_artifact_format(self) -> None:
        with self.assertRaises(UnsupportedEvidenceFormat):
            build_evidence_corpus([EvidenceSource(artifact_name="image.png", data=b"PNG")])

    def test_validates_quote_within_requested_reference_after_normalization(self) -> None:
        corpus = self._sample_corpus()
        result = validate_evidence_quote(
            corpus,
            "“人工智能系统 能够提升城市治理效率”",
            reference_id=1,
            min_quote_chars=8,
        )

        self.assertTrue(result.valid)
        self.assertEqual("matched", result.status)
        self.assertEqual(1, result.resolved_reference_id)
        self.assertEqual(1, len(result.matches))

    def test_reports_when_quote_only_exists_under_another_reference(self) -> None:
        corpus = self._sample_corpus()
        result = validate_evidence_quote(
            corpus,
            "第二篇研究发现部署成本仍然较高",
            reference_id=1,
            min_quote_chars=8,
        )

        self.assertFalse(result.valid)
        self.assertEqual("reference_mismatch", result.status)
        self.assertEqual(2, result.matches[0].reference_id)

    def test_global_match_requires_one_unambiguous_chunk(self) -> None:
        unique = self._sample_corpus()
        matched = validate_evidence_quote(unique, "第二篇研究发现部署成本仍然较高", min_quote_chars=8)
        self.assertTrue(matched.valid)
        self.assertEqual(2, matched.resolved_reference_id)

        ambiguous = build_evidence_corpus(
            [
                EvidenceSource(artifact_name="[1] a.txt", data="相同的证据原句具有足够长度。".encode()),
                EvidenceSource(artifact_name="[2] b.txt", data="相同的证据原句具有足够长度。".encode()),
            ]
        )
        result = validate_evidence_quote(ambiguous, "相同的证据原句具有足够长度", min_quote_chars=8)
        self.assertFalse(result.valid)
        self.assertEqual("ambiguous", result.status)
        self.assertEqual(2, len(result.matches))

    def test_rejects_short_and_unknown_quotes(self) -> None:
        corpus = self._sample_corpus()
        short = validate_evidence_quote(corpus, "效率", reference_id=1)
        missing_reference = validate_evidence_quote(
            corpus,
            "这是一条长度充分但并不存在的原文引句",
            reference_id=9,
            min_quote_chars=8,
        )
        unknown = validate_evidence_quote(
            corpus,
            "这是一条长度充分但并不存在的原文引句",
            reference_id=1,
            min_quote_chars=8,
        )

        self.assertEqual("quote_too_short", short.status)
        self.assertEqual("reference_not_found", missing_reference.status)
        self.assertEqual("not_found", unknown.status)

    def test_audit_requires_complete_occurrence_coverage_and_exact_chunk_match(self) -> None:
        corpus = self._sample_corpus()
        ref1, ref2 = corpus.blocks
        occurrences = [
            CitationOccurrence("section-1:citation-1", 1),
            CitationOccurrence("section-2:citation-1", 2),
        ]
        assertions = [
            EvidenceAssertion(
                "section-1:citation-1",
                ref1.chunk_id,
                "人工智能系统能够提升城市治理效率",
            ),
            EvidenceAssertion(
                "section-2:citation-1",
                ref2.chunk_id,
                "第二篇研究发现部署成本仍然较高",
            ),
        ]

        result = verify_evidence_audit(corpus, occurrences, assertions, min_quote_chars=8)

        self.assertTrue(result.valid)
        self.assertEqual(2, result.matched_count)
        self.assertTrue(all(check.locator is not None for check in result.checks))
        self.assertEqual(2, result.to_dict()["occurrence_count"])

    def test_audit_reports_missing_unknown_duplicate_and_unmatched_assertions(self) -> None:
        corpus = self._sample_corpus()
        ref1, ref2 = corpus.blocks
        occurrences = [
            CitationOccurrence("occ-1", 1),
            CitationOccurrence("occ-2", 2),
            CitationOccurrence("occ-3", 1),
            CitationOccurrence("occ-4", 1),
            CitationOccurrence("occ-5", 1),
        ]
        assertions = [
            EvidenceAssertion("occ-2", ref1.chunk_id, "人工智能系统能够提升城市治理效率"),
            EvidenceAssertion("occ-3", "missing-chunk", "足够长但找不到对应块的证据原句"),
            EvidenceAssertion("occ-4", ref1.chunk_id, "足够长但不在原始证据块中的虚构引句"),
            EvidenceAssertion("occ-5", ref1.chunk_id, "人工智能系统能够提升城市治理效率"),
            EvidenceAssertion("occ-5", ref2.chunk_id, "第二篇研究发现部署成本仍然较高"),
            EvidenceAssertion("unknown-occurrence", ref1.chunk_id, "人工智能系统能够提升城市治理效率"),
        ]

        result = verify_evidence_audit(corpus, occurrences, assertions, min_quote_chars=8)

        self.assertFalse(result.valid)
        self.assertEqual(("occ-1",), result.missing_occurrence_ids)
        self.assertEqual(("unknown-occurrence",), result.unknown_occurrence_ids)
        self.assertEqual(("occ-5",), result.duplicate_occurrence_ids)
        self.assertEqual(("missing-chunk",), result.unknown_chunk_ids)
        self.assertEqual(("occ-2", "occ-3", "occ-4"), result.unmatched_occurrence_ids)
        self.assertEqual(
            [
                "missing_assertion",
                "reference_mismatch",
                "unknown_chunk",
                "quote_not_found",
                "duplicate_assertion",
            ],
            [check.status for check in result.checks],
        )

    @staticmethod
    def _sample_corpus():
        return build_evidence_corpus(
            [
                EvidenceSource(
                    artifact_name="[1] governance.txt",
                    artifact_id=1,
                    data="人工智能系统能够提升城市治理效率。".encode(),
                ),
                EvidenceSource(
                    artifact_name="ref-2-cost.txt",
                    artifact_id=2,
                    data="第二篇研究发现部署成本仍然较高。".encode(),
                ),
            ]
        )


if __name__ == "__main__":
    unittest.main()
