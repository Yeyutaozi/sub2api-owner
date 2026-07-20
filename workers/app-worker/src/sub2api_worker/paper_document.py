"""Build standards-friendly academic paper DOCX artifacts.

The module deliberately keeps content generation separate from document
rendering. Callers provide structured paper content and formatting settings;
the returned bytes can be uploaded directly as a Worker artifact.
"""

from __future__ import annotations

import copy
import io
import re
from collections.abc import Mapping, Sequence
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


PAGE_SIZES_MM: dict[str, tuple[float, float]] = {
    "A3": (297.0, 420.0),
    "A4": (210.0, 297.0),
    "A5": (148.0, 210.0),
    "B5": (176.0, 250.0),
    "B5_JIS": (182.0, 257.0),
    "LETTER": (215.9, 279.4),
    "LEGAL": (215.9, 355.6),
}


DEFAULT_ACADEMIC_PAPER_SETTINGS: dict[str, Any] = {
    "format_preset": "academic_cn",
    "citation_style": "",
    "page": {
        "size": "A4",
        "orientation": "portrait",
        "width_mm": 210.0,
        "height_mm": 297.0,
        "margins_mm": {
            "top": 25.4,
            "bottom": 25.4,
            "left": 30.0,
            "right": 25.4,
            "gutter": 0.0,
        },
        "header_distance_mm": 15.0,
        "footer_distance_mm": 15.0,
    },
    "fonts": {
        "east_asia": "宋体",
        "latin": "Times New Roman",
        "heading_east_asia": "黑体",
        "heading_latin": "Times New Roman",
    },
    "title": {
        "east_asia_font": "黑体",
        "latin_font": "Times New Roman",
        "size_pt": 22.0,
        "bold": True,
        "alignment": "center",
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.25,
        "space_before_pt": 0.0,
        "space_after_pt": 18.0,
    },
    "subtitle": {
        "east_asia_font": "黑体",
        "latin_font": "Times New Roman",
        "size_pt": 16.0,
        "bold": True,
        "alignment": "center",
        "space_before_pt": 0.0,
        "space_after_pt": 16.0,
    },
    "author": {
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 12.0,
        "alignment": "center",
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.5,
        "space_before_pt": 0.0,
        "space_after_pt": 6.0,
    },
    "metadata": {
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 12.0,
        "alignment": "center",
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.5,
        "space_before_pt": 0.0,
        "space_after_pt": 4.0,
    },
    "body": {
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 12.0,
        "bold": False,
        "italic": False,
        "alignment": "justify",
        "first_line_indent_chars": 2.0,
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.5,
        "space_before_pt": 0.0,
        "space_after_pt": 0.0,
        "widow_control": True,
    },
    "abstract": {
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 12.0,
        "alignment": "justify",
        "first_line_indent_chars": 2.0,
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.5,
        "space_before_pt": 0.0,
        "space_after_pt": 0.0,
        "widow_control": True,
    },
    "keywords": {
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 12.0,
        "alignment": "left",
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.5,
        "space_before_pt": 6.0,
        "space_after_pt": 0.0,
    },
    "references": {
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 10.5,
        "alignment": "justify",
        "first_line_indent_chars": 0.0,
        "left_indent_cm": 0.74,
        "hanging_indent_cm": 0.74,
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.25,
        "space_before_pt": 0.0,
        "space_after_pt": 3.0,
    },
    "acknowledgements": {
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 12.0,
        "alignment": "justify",
        "first_line_indent_chars": 2.0,
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.5,
        "space_before_pt": 0.0,
        "space_after_pt": 0.0,
    },
    "appendix": {
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 12.0,
        "alignment": "justify",
        "first_line_indent_chars": 2.0,
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.5,
        "space_before_pt": 0.0,
        "space_after_pt": 0.0,
    },
    "headings": {
        "1": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 16.0,
            "bold": True,
            "alignment": "center",
            "line_spacing_mode": "multiple",
            "line_spacing_value": 1.5,
            "space_before_pt": 18.0,
            "space_after_pt": 12.0,
            "keep_with_next": True,
            "widow_control": True,
        },
        "2": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 15.0,
            "bold": True,
            "alignment": "left",
            "line_spacing_mode": "multiple",
            "line_spacing_value": 1.5,
            "space_before_pt": 12.0,
            "space_after_pt": 6.0,
            "keep_with_next": True,
            "widow_control": True,
        },
        "3": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 14.0,
            "bold": True,
            "alignment": "left",
            "line_spacing_mode": "multiple",
            "line_spacing_value": 1.5,
            "space_before_pt": 9.0,
            "space_after_pt": 6.0,
            "keep_with_next": True,
            "widow_control": True,
        },
        "4": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 12.0,
            "bold": True,
            "alignment": "left",
            "line_spacing_mode": "multiple",
            "line_spacing_value": 1.5,
            "space_before_pt": 6.0,
            "space_after_pt": 3.0,
            "keep_with_next": True,
            "widow_control": True,
        },
        "5": {
            "east_asia_font": "宋体",
            "latin_font": "Times New Roman",
            "size_pt": 12.0,
            "bold": True,
            "alignment": "left",
            "line_spacing_mode": "multiple",
            "line_spacing_value": 1.5,
            "space_before_pt": 6.0,
            "space_after_pt": 3.0,
            "keep_with_next": True,
            "widow_control": True,
        },
    },
    "heading_numbering": {
        "enabled": True,
        "formats": ["第%1章", "%1.%2", "%1.%2.%3", "%1.%2.%3.%4", "%1.%2.%3.%4.%5"],
        "number_formats": ["decimal", "decimal", "decimal", "decimal", "decimal"],
        "starts": [1, 1, 1, 1, 1],
        "suffix": "space",
        "strip_existing_numbers": True,
    },
    "toc": {
        "enabled": True,
        "title": "目录",
        "levels": 5,
        "before_abstract": False,
        "page_break_before": True,
        "page_break_after": True,
    },
    "header": {
        "enabled": False,
        "text": "",
        "use_title": False,
        "alignment": "center",
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 9.0,
        "bold": False,
        "italic": False,
        "different_first_page": True,
        "distance_mm": 15.0,
    },
    "footer": {
        "enabled": False,
        "text": "",
        "alignment": "center",
        "east_asia_font": "宋体",
        "latin_font": "Times New Roman",
        "size_pt": 9.0,
        "bold": False,
        "italic": False,
        "different_first_page": True,
        "distance_mm": 15.0,
    },
    "page_number": {
        "enabled": True,
        "position": "footer",
        "alignment": "center",
        "start": 1,
        "format": "decimal",
        "show_on_first_page": False,
        "prefix": "",
        "suffix": "",
    },
    "pagination": {
        "include_title_page": True,
        "title_page_break_after": True,
        "abstract_page_break_after": True,
        "toc_page_break_after": True,
        "chapter_page_break_before": True,
        "references_page_break_before": True,
        "acknowledgements_page_break_before": True,
        "appendix_page_break_before": True,
    },
    "labels": {
        "abstract": "摘要",
        "keywords": "关键词",
        "references": "参考文献",
        "acknowledgements": "致谢",
        "appendix": "附录",
    },
}


ACADEMIC_FORMAT_PRESETS: dict[str, dict[str, Any]] = {
    "academic_cn": {},
    "standard_cn_academic": {},
    "undergraduate_thesis": {
        "page": {"size": "A4", "orientation": "portrait"},
        "body": {"size_pt": 12.0, "line_spacing_mode": "multiple", "line_spacing_value": 1.5},
        "headings": {
            "1": {"east_asia_font": "黑体", "size_pt": 16.0, "bold": True, "alignment": "center"},
            "2": {"east_asia_font": "黑体", "size_pt": 15.0, "bold": True, "alignment": "left"},
            "3": {"east_asia_font": "黑体", "size_pt": 14.0, "bold": True, "alignment": "left"},
        },
    },
    "graduate_thesis": {
        "page": {"size": "A4", "orientation": "portrait", "margins_mm": {"left": 30.0, "right": 25.0}},
        "body": {"size_pt": 12.0, "line_spacing_mode": "exact", "line_spacing_value": 20.0},
        "abstract": {"size_pt": 12.0, "line_spacing_mode": "exact", "line_spacing_value": 20.0},
        "headings": {
            "1": {"east_asia_font": "黑体", "size_pt": 16.0, "bold": True, "alignment": "center"},
            "2": {"east_asia_font": "黑体", "size_pt": 14.0, "bold": True, "alignment": "left"},
            "3": {"east_asia_font": "黑体", "size_pt": 12.0, "bold": True, "alignment": "left"},
            "4": {"east_asia_font": "宋体", "size_pt": 12.0, "bold": True, "alignment": "left"},
            "5": {"east_asia_font": "宋体", "size_pt": 12.0, "bold": False, "alignment": "left"},
        },
    },
    "cn_journal": {
        "page": {"size": "A4", "orientation": "portrait", "margins_mm": {"top": 20.0, "bottom": 20.0, "left": 20.0, "right": 20.0}},
        "title": {"size_pt": 18.0, "space_after_pt": 10.0},
        "body": {"size_pt": 10.5, "line_spacing_mode": "multiple", "line_spacing_value": 1.2, "space_after_pt": 3.0},
        "abstract": {"size_pt": 9.0, "line_spacing_mode": "multiple", "line_spacing_value": 1.15},
        "keywords": {"size_pt": 9.0, "line_spacing_mode": "multiple", "line_spacing_value": 1.15},
        "headings": {
            "1": {"size_pt": 12.0, "alignment": "left", "space_before_pt": 9.0, "space_after_pt": 4.0},
            "2": {"size_pt": 10.5, "alignment": "left", "space_before_pt": 6.0, "space_after_pt": 3.0},
            "3": {"size_pt": 10.5, "alignment": "left", "space_before_pt": 4.0, "space_after_pt": 2.0},
            "4": {"size_pt": 9.0, "alignment": "left"},
            "5": {"size_pt": 9.0, "alignment": "left"},
        },
        "toc": {"enabled": False},
        "pagination": {"include_title_page": False, "chapter_page_break_before": False},
    },
    "apa_english": {
        "page": {"size": "LETTER", "orientation": "portrait", "margins_mm": {"top": 25.4, "bottom": 25.4, "left": 25.4, "right": 25.4}},
        "fonts": {"east_asia": "Times New Roman", "latin": "Times New Roman", "heading_east_asia": "Times New Roman", "heading_latin": "Times New Roman"},
        "title": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "bold": True, "alignment": "center"},
        "subtitle": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "bold": False, "alignment": "center"},
        "author": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "alignment": "center"},
        "metadata": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "alignment": "center"},
        "body": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "alignment": "left", "first_line_indent_cm": 1.27, "first_line_indent_chars": None, "line_spacing_mode": "multiple", "line_spacing_value": 2.0, "space_after_pt": 0.0},
        "abstract": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "alignment": "left", "first_line_indent_cm": 1.27, "first_line_indent_chars": None, "line_spacing_mode": "multiple", "line_spacing_value": 2.0},
        "keywords": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "line_spacing_mode": "multiple", "line_spacing_value": 2.0},
        "references": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "line_spacing_mode": "multiple", "line_spacing_value": 2.0},
        "acknowledgements": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "line_spacing_mode": "multiple", "line_spacing_value": 2.0},
        "appendix": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "line_spacing_mode": "multiple", "line_spacing_value": 2.0},
        "headings": {
            "1": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "bold": True, "italic": False, "alignment": "center"},
            "2": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "bold": True, "italic": False, "alignment": "left"},
            "3": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "bold": True, "italic": True, "alignment": "left"},
            "4": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "bold": True, "italic": False, "alignment": "left"},
            "5": {"east_asia_font": "Times New Roman", "latin_font": "Times New Roman", "size_pt": 12.0, "bold": False, "italic": True, "alignment": "left"},
        },
        "heading_numbering": {"enabled": False},
    },
    "uploaded_template": {},
    "custom": {},
}


_ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "centre": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distributed": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}

_STYLE_SUFFIXES = {
    "east_asia_font",
    "latin_font",
    "size_pt",
    "bold",
    "italic",
    "underline",
    "color",
    "alignment",
    "first_line_indent_chars",
    "first_line_indent_cm",
    "left_indent_cm",
    "right_indent_cm",
    "hanging_indent_cm",
    "line_spacing_mode",
    "line_spacing_value",
    "space_before_pt",
    "space_after_pt",
    "keep_with_next",
    "keep_together",
    "page_break_before",
    "widow_control",
    "letter_spacing_pt",
    "font_scale_percent",
}

_CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
_LATIN_WORD_RE = re.compile(r"[A-Za-z]+(?:['-][A-Za-z]+)*|\d+(?:[.,]\d+)*")
_CODE_FENCE_RE = re.compile(r"^\s*```.*?$", re.MULTILINE)
_HEADING_MARKER_RE = re.compile(r"^\s{0,3}#{1,6}\s+", re.MULTILINE)
_INLINE_LINK_RE = re.compile(r"!?\[([^\]]*)\]\([^)]*\)")
_LEADING_LIST_RE = re.compile(r"^\s*(?:[-*+]\s+|\d+[.)、]\s+)")
_EXISTING_HEADING_NUMBER_RE = re.compile(
    r"^\s*(?:第[一二三四五六七八九十百千万\d]+[章节篇]|\d+(?:\.\d+)*[.)、]?|[一二三四五六七八九十]+[、.])\s+"
)


def _as_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return value != 0
    return str(value).strip().lower() in {"1", "true", "yes", "on", "enabled", "是"}


def _as_float(value: Any, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _as_int(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _deep_merge(base: Mapping[str, Any], override: Mapping[str, Any]) -> dict[str, Any]:
    result = copy.deepcopy(dict(base))
    for key, value in override.items():
        if isinstance(value, Mapping) and isinstance(result.get(key), Mapping):
            result[key] = _deep_merge(result[key], value)
        else:
            result[key] = copy.deepcopy(value)
    return result


def _set_path(target: dict[str, Any], path: Sequence[str], value: Any) -> None:
    cursor = target
    for key in path[:-1]:
        child = cursor.get(key)
        if not isinstance(child, dict):
            child = {}
            cursor[key] = child
        cursor = child
    cursor[path[-1]] = value


def _normalise_settings(settings: Mapping[str, Any] | None) -> dict[str, Any]:
    incoming = copy.deepcopy(dict(settings or {}))
    aliases: dict[str, Any] = {}

    direct_paths = {
        "format_preset": ("format_preset",),
        "page_size": ("page", "size"),
        "orientation": ("page", "orientation"),
        "page_orientation": ("page", "orientation"),
        "page_width_mm": ("page", "width_mm"),
        "page_height_mm": ("page", "height_mm"),
        "page_header_distance_mm": ("page", "header_distance_mm"),
        "page_footer_distance_mm": ("page", "footer_distance_mm"),
        "cover_enabled": ("pagination", "include_title_page"),
        "cover_page_break_after": ("pagination", "title_page_break_after"),
    }
    for flat_key, path in direct_paths.items():
        if flat_key in incoming:
            _set_path(aliases, path, incoming[flat_key])
    for flat_key in ("default_east_asia_font", "east_asia_font"):
        if flat_key in incoming:
            _set_path(aliases, ("fonts", "east_asia"), incoming[flat_key])
            break
    for flat_key in ("default_latin_font", "latin_font"):
        if flat_key in incoming:
            _set_path(aliases, ("fonts", "latin"), incoming[flat_key])
            break

    for side in ("top", "bottom", "left", "right", "gutter"):
        cm_key = f"page_margin_{side}_cm"
        mm_key = f"page_margin_{side}_mm"
        if cm_key in incoming:
            _set_path(aliases, ("page", "margins_mm", side), _as_float(incoming[cm_key], 0.0) * 10.0)
        elif mm_key in incoming:
            _set_path(aliases, ("page", "margins_mm", side), incoming[mm_key])

    style_prefixes = {
        "cover_title_": "title",
        "cover_subtitle_": "subtitle",
        "cover_author_": "author",
        "cover_metadata_": "metadata",
        "title_": "title",
        "subtitle_": "subtitle",
        "author_": "author",
        "metadata_": "metadata",
        "body_": "body",
        "abstract_": "abstract",
        "keywords_": "keywords",
        "references_": "references",
        "acknowledgements_": "acknowledgements",
        "appendix_": "appendix",
    }
    for prefix, section_name in style_prefixes.items():
        for suffix in _STYLE_SUFFIXES:
            key = f"{prefix}{suffix}"
            if key in incoming:
                _set_path(aliases, (section_name, suffix), incoming[key])

    for level in range(1, 6):
        for suffix in _STYLE_SUFFIXES:
            key = f"heading{level}_{suffix}"
            if key in incoming:
                _set_path(aliases, ("headings", str(level), suffix), incoming[key])

    for group in ("toc", "header", "footer", "page_number", "pagination", "heading_numbering"):
        prefix = f"{group}_"
        for key, value in incoming.items():
            if key.startswith(prefix) and len(key) > len(prefix):
                _set_path(aliases, (group, key[len(prefix) :]), value)

    nested = {
        key: value
        for key, value in incoming.items()
        if isinstance(value, Mapping)
        or key in DEFAULT_ACADEMIC_PAPER_SETTINGS
        or key == "format_preset"
    }
    explicit = _deep_merge(aliases, nested)
    preset_name = str(explicit.get("format_preset", DEFAULT_ACADEMIC_PAPER_SETTINGS["format_preset"])).strip().lower()
    if preset_name not in ACADEMIC_FORMAT_PRESETS:
        supported = ", ".join(ACADEMIC_FORMAT_PRESETS)
        raise ValueError(f"Unsupported format_preset {preset_name!r}; expected one of: {supported}")
    preset_base = _deep_merge(DEFAULT_ACADEMIC_PAPER_SETTINGS, ACADEMIC_FORMAT_PRESETS[preset_name])
    resolved = _deep_merge(preset_base, explicit)
    explicit_fonts = explicit.get("fonts", {}) if isinstance(explicit, Mapping) else {}
    if isinstance(explicit_fonts, Mapping):
        style_names = (
            "title",
            "subtitle",
            "author",
            "metadata",
            "body",
            "abstract",
            "keywords",
            "references",
            "acknowledgements",
            "appendix",
            "header",
            "footer",
        )
        for font_key, style_key in (("east_asia", "east_asia_font"), ("latin", "latin_font")):
            if font_key not in explicit_fonts:
                continue
            for style_name in style_names:
                explicit_style = explicit.get(style_name, {}) if isinstance(explicit, Mapping) else {}
                if not isinstance(explicit_style, Mapping) or style_key not in explicit_style:
                    resolved[style_name][style_key] = explicit_fonts[font_key]
            explicit_headings = explicit.get("headings", {}) if isinstance(explicit, Mapping) else {}
            for level in range(1, 6):
                explicit_heading = (
                    explicit_headings.get(str(level), {}) if isinstance(explicit_headings, Mapping) else {}
                )
                if not isinstance(explicit_heading, Mapping) or style_key not in explicit_heading:
                    resolved["headings"][str(level)][style_key] = explicit_fonts[font_key]
    resolved["format_preset"] = preset_name
    resolved["_explicit_settings"] = explicit
    resolved["_preserve_template_styles"] = preset_name == "uploaded_template"
    return resolved


def _clean_markdown_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    text = _CODE_FENCE_RE.sub("", text)
    text = _HEADING_MARKER_RE.sub("", text)
    text = _INLINE_LINK_RE.sub(r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def _clean_heading_text(value: Any, strip_number: bool) -> str:
    text = _clean_markdown_text(value)
    if strip_number:
        text = _EXISTING_HEADING_NUMBER_RE.sub("", text)
    return text.strip()


def _join_wrapped_lines(lines: Sequence[str]) -> str:
    result = ""
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            continue
        if not result:
            result = cleaned
            continue
        if _CJK_RE.search(result[-1:]) or _CJK_RE.search(cleaned[:1]):
            result += cleaned
        else:
            result += " " + cleaned
    return result


def _paragraph_blocks(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, Mapping):
        for key in ("text", "content", "paragraphs", "value"):
            if key in value:
                return _paragraph_blocks(value[key])
        return []
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
        blocks: list[str] = []
        for item in value:
            blocks.extend(_paragraph_blocks(item))
        return blocks

    raw = str(value).replace("\r\n", "\n").replace("\r", "\n")
    raw = _CODE_FENCE_RE.sub("", raw)
    groups = re.split(r"\n\s*\n+", raw)
    blocks = []
    for group in groups:
        lines = [line for line in group.split("\n") if line.strip()]
        if not lines:
            continue
        list_like = len(lines) > 1 and all(_LEADING_LIST_RE.match(line) for line in lines)
        if list_like:
            candidates = [_LEADING_LIST_RE.sub("", line) for line in lines]
        else:
            candidates = [_join_wrapped_lines(lines)]
        for candidate in candidates:
            cleaned = _clean_markdown_text(candidate)
            if cleaned:
                blocks.append(cleaned)
    return blocks


def _text_from_reference(reference: Any) -> str:
    if isinstance(reference, Mapping):
        for key in ("citation", "text", "formatted", "content", "title"):
            value = reference.get(key)
            if value:
                return _clean_markdown_text(value)
        return ""
    return _clean_markdown_text(reference)


_REFERENCE_NUMBER_PREFIX_RE = re.compile(r"^\s*(?:\[\s*\d+\s*\]|【\s*\d+\s*】)\s*")


def _uses_numeric_reference_style(settings: Mapping[str, Any]) -> bool:
    style = str(settings.get("citation_style", "")).strip().lower().replace("-", "_").replace(" ", "")
    if style in {
        "gbt7714_numeric",
        "gb/t7714_numeric",
        "numeric",
        "numbered",
        "ieee",
        "vancouver",
    }:
        return True
    if "7714" in style and not any(marker in style for marker in ("author", "year", "著者", "年份")):
        return True
    return False


def _numbered_reference_text(reference: Any, index: int) -> str:
    text = _REFERENCE_NUMBER_PREFIX_RE.sub("", _text_from_reference(reference), count=1).strip()
    return f"[{index}] {text}" if text else ""


def _iter_sections(sections: Any, inferred_level: int = 1):
    if isinstance(sections, Mapping):
        sections = [sections]
    if not isinstance(sections, Sequence) or isinstance(sections, (str, bytes, bytearray)):
        return
    for item in sections:
        if isinstance(item, str):
            yield {"title": item, "level": inferred_level, "content": "", "children": []}
            continue
        if not isinstance(item, Mapping):
            continue
        level = max(1, min(5, _as_int(item.get("level"), inferred_level)))
        title = item.get("title") or item.get("heading") or item.get("name") or ""
        content = item.get("content", item.get("text", item.get("paragraphs", "")))
        children = item.get("children", item.get("subsections", item.get("sections", [])))
        yield {"title": title, "level": level, "content": content, "children": children}


def _all_sections(sections: Any, inferred_level: int = 1):
    for section in _iter_sections(sections, inferred_level) or ():
        yield section
        yield from _all_sections(section["children"], min(5, section["level"] + 1))


def _value_text(value: Any, separator: str = "、") -> str:
    if value is None:
        return ""
    if isinstance(value, Mapping):
        return separator.join(_clean_markdown_text(item) for item in value.values() if item not in (None, ""))
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
        return separator.join(_clean_markdown_text(item) for item in value if item not in (None, ""))
    return _clean_markdown_text(value)


def _paper_text_parts(paper: Mapping[str, Any], include_front_matter: bool = True) -> list[str]:
    parts: list[str] = []
    if include_front_matter:
        for key in ("title", "subtitle", "abstract"):
            parts.extend(_paragraph_blocks(paper.get(key)))
        keywords = _value_text(paper.get("keywords"))
        if keywords:
            parts.append(keywords)
    for section in _all_sections(paper.get("sections", paper.get("outline", []))):
        title = _clean_markdown_text(section["title"])
        if title:
            parts.append(title)
        parts.extend(_paragraph_blocks(section["content"]))
    references = paper.get("references", [])
    if isinstance(references, Sequence) and not isinstance(references, (str, bytes, bytearray)):
        parts.extend(filter(None, (_text_from_reference(item) for item in references)))
    else:
        parts.extend(_paragraph_blocks(references))
    parts.extend(_paragraph_blocks(paper.get("acknowledgements", paper.get("thanks"))))
    appendices = paper.get("appendices", paper.get("appendix", []))
    if isinstance(appendices, Mapping):
        appendices = [appendices]
    if isinstance(appendices, Sequence) and not isinstance(appendices, (str, bytes, bytearray)):
        for appendix in appendices:
            if isinstance(appendix, Mapping):
                parts.extend(_paragraph_blocks(appendix.get("title")))
                parts.extend(_paragraph_blocks(appendix.get("content", appendix.get("text"))))
            else:
                parts.extend(_paragraph_blocks(appendix))
    return [part for part in parts if part]


def _text_metrics(parts: Sequence[str]) -> dict[str, int]:
    text = "\n".join(parts)
    clean = _clean_markdown_text(text)
    cjk = len(_CJK_RE.findall(clean))
    latin_words = len(_LATIN_WORD_RE.findall(clean))
    return {
        "text_units": cjk + latin_words,
        "cjk_characters": cjk,
        "latin_words": latin_words,
        "characters_no_whitespace": len(re.sub(r"\s+", "", clean)),
        "paragraphs": len([part for part in parts if part.strip()]),
    }


def count_academic_paper_words(paper: Mapping[str, Any]) -> dict[str, int]:
    """Return Chinese-character/Western-word counts for a paper."""

    total_parts = _paper_text_parts(paper, include_front_matter=True)
    body_parts = _paper_text_parts(paper, include_front_matter=False)
    total = _text_metrics(total_parts)
    body = _text_metrics(body_parts)
    section_count = sum(1 for _ in _all_sections(paper.get("sections", paper.get("outline", []))))
    return {
        "word_count": total["text_units"],
        "body_word_count": body["text_units"],
        "cjk_characters": total["cjk_characters"],
        "latin_words": total["latin_words"],
        "characters_no_whitespace": total["characters_no_whitespace"],
        "paragraph_count": total["paragraphs"],
        "section_count": section_count,
    }


def _page_dimensions(settings: Mapping[str, Any]) -> tuple[float, float, str, str]:
    page = settings["page"]
    size = str(page.get("size", "A4")).strip().upper().replace("-", "_")
    if size in {"US_LETTER", "LETTER_SIZE"}:
        size = "LETTER"
    if size in {"US_LEGAL", "LEGAL_SIZE"}:
        size = "LEGAL"
    if size in {"JIS_B5", "B5JIS"}:
        size = "B5_JIS"
    if size == "CUSTOM":
        width = _as_float(page.get("width_mm"), 0.0)
        height = _as_float(page.get("height_mm"), 0.0)
        if width <= 0 or height <= 0:
            raise ValueError("CUSTOM page size requires positive page.width_mm and page.height_mm")
    elif size in PAGE_SIZES_MM:
        width, height = PAGE_SIZES_MM[size]
    else:
        supported = ", ".join([*PAGE_SIZES_MM, "CUSTOM"])
        raise ValueError(f"Unsupported page size {size!r}; expected one of: {supported}")
    orientation = str(page.get("orientation", "portrait")).strip().lower()
    if orientation not in {"portrait", "landscape"}:
        raise ValueError("page.orientation must be 'portrait' or 'landscape'")
    if orientation == "landscape":
        width, height = max(width, height), min(width, height)
    else:
        width, height = min(width, height), max(width, height)
    return width, height, size, orientation


def _line_spacing_summary(style: Mapping[str, Any]) -> str:
    mode = str(style.get("line_spacing_mode", "single")).lower()
    value = style.get("line_spacing_value")
    if mode in {"exact", "fixed", "at_least"}:
        return f"{mode}:{_as_float(value, 12.0):g}pt"
    if mode == "double":
        return "multiple:2"
    if mode == "one_and_half":
        return "multiple:1.5"
    if mode == "single":
        return "single"
    return f"multiple:{_as_float(value, 1.0):g}"


def summarize_academic_paper_format(settings: Mapping[str, Any] | None = None) -> dict[str, Any]:
    """Return a compact, JSON-safe summary of the resolved DOCX format."""

    resolved = _normalise_settings(settings)
    width, height, size, orientation = _page_dimensions(resolved)
    body = resolved["body"]
    headings = {
        str(level): {
            "east_asia_font": resolved["headings"][str(level)].get("east_asia_font"),
            "latin_font": resolved["headings"][str(level)].get("latin_font"),
            "size_pt": _as_float(resolved["headings"][str(level)].get("size_pt"), 12.0),
            "bold": _as_bool(resolved["headings"][str(level)].get("bold"), True),
            "alignment": resolved["headings"][str(level)].get("alignment", "left"),
            "line_spacing": _line_spacing_summary(resolved["headings"][str(level)]),
        }
        for level in range(1, 6)
    }
    return {
        "format_preset": resolved.get("format_preset", "academic_cn"),
        "page": {
            "size": size,
            "orientation": orientation,
            "width_mm": round(width, 2),
            "height_mm": round(height, 2),
            "margins_mm": copy.deepcopy(resolved["page"]["margins_mm"]),
        },
        "body": {
            "east_asia_font": body.get("east_asia_font"),
            "latin_font": body.get("latin_font"),
            "size_pt": _as_float(body.get("size_pt"), 12.0),
            "alignment": body.get("alignment", "justify"),
            "first_line_indent_chars": _as_float(body.get("first_line_indent_chars"), 0.0),
            "first_line_indent_cm": body.get("first_line_indent_cm"),
            "line_spacing": _line_spacing_summary(body),
        },
        "headings": headings,
        "heading_numbering": {
            "enabled": _as_bool(resolved["heading_numbering"].get("enabled"), True),
            "formats": list(resolved["heading_numbering"].get("formats", []))[:5],
        },
        "toc": {
            "enabled": _as_bool(resolved["toc"].get("enabled"), True),
            "levels": max(1, min(5, _as_int(resolved["toc"].get("levels"), 5))),
        },
        "page_number": {
            "enabled": _as_bool(resolved["page_number"].get("enabled"), True),
            "position": resolved["page_number"].get("position", "footer"),
            "start": _as_int(resolved["page_number"].get("start"), 1),
            "format": resolved["page_number"].get("format", "decimal"),
        },
    }


def summarize_academic_paper(
    paper: Mapping[str, Any], settings: Mapping[str, Any] | None = None
) -> dict[str, Any]:
    """Return content statistics and the resolved formatting summary."""

    return {
        "title": _clean_markdown_text(paper.get("title", "")),
        "content": count_academic_paper_words(paper),
        "format": summarize_academic_paper_format(settings),
    }


def _set_rfonts(r_pr: Any, east_asia: str, latin: str) -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)


def _set_or_remove_child(parent: Any, tag: str, attr: str, value: str | None) -> None:
    child = parent.find(qn(tag))
    if value is None:
        if child is not None:
            parent.remove(child)
        return
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    child.set(qn(attr), value)


def _apply_run_style(style: Any, config: Mapping[str, Any]) -> None:
    east_asia = str(config.get("east_asia_font") or "宋体")
    latin = str(config.get("latin_font") or "Times New Roman")
    style.font.name = latin
    style.font.size = Pt(max(1.0, _as_float(config.get("size_pt"), 12.0)))
    style.font.bold = _as_bool(config.get("bold"), False)
    style.font.italic = _as_bool(config.get("italic"), False)
    style.font.underline = _as_bool(config.get("underline"), False)
    color = str(config.get("color", "000000")).strip().lstrip("#")
    if re.fullmatch(r"[0-9A-Fa-f]{6}", color):
        style.font.color.rgb = RGBColor.from_string(color.upper())
    r_pr = style.element.get_or_add_rPr()
    _set_rfonts(r_pr, east_asia, latin)
    if "letter_spacing_pt" in config:
        spacing = round(_as_float(config.get("letter_spacing_pt"), 0.0) * 20)
        _set_or_remove_child(r_pr, "w:spacing", "w:val", str(spacing))
    if "font_scale_percent" in config:
        scale = max(1, min(600, _as_int(config.get("font_scale_percent"), 100)))
        _set_or_remove_child(r_pr, "w:w", "w:val", str(scale))


def _apply_paragraph_style(style: Any, config: Mapping[str, Any], outline_level: int | None = None) -> None:
    _apply_run_style(style, config)
    paragraph = style.paragraph_format
    alignment = str(config.get("alignment", "left")).strip().lower()
    paragraph.alignment = _ALIGNMENTS.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.space_before = Pt(max(0.0, _as_float(config.get("space_before_pt"), 0.0)))
    paragraph.space_after = Pt(max(0.0, _as_float(config.get("space_after_pt"), 0.0)))
    paragraph.keep_with_next = _as_bool(config.get("keep_with_next"), False)
    paragraph.keep_together = _as_bool(config.get("keep_together"), False)
    paragraph.page_break_before = _as_bool(config.get("page_break_before"), False)
    paragraph.widow_control = _as_bool(config.get("widow_control"), True)
    if config.get("left_indent_cm") is not None:
        paragraph.left_indent = Cm(_as_float(config.get("left_indent_cm"), 0.0))
    if config.get("right_indent_cm") is not None:
        paragraph.right_indent = Cm(_as_float(config.get("right_indent_cm"), 0.0))
    if config.get("hanging_indent_cm") is not None:
        paragraph.first_line_indent = Cm(-abs(_as_float(config.get("hanging_indent_cm"), 0.0)))
    elif config.get("first_line_indent_cm") is not None:
        paragraph.first_line_indent = Cm(_as_float(config.get("first_line_indent_cm"), 0.0))
    elif config.get("first_line_indent_chars") is not None:
        chars = _as_float(config.get("first_line_indent_chars"), 0.0)
        size = _as_float(config.get("size_pt"), 12.0)
        paragraph.first_line_indent = Pt(chars * size)

    mode = str(config.get("line_spacing_mode", "single")).strip().lower()
    value = _as_float(config.get("line_spacing_value"), 1.0)
    if mode in {"exact", "fixed"}:
        paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.line_spacing = Pt(max(1.0, value))
    elif mode == "at_least":
        paragraph.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        paragraph.line_spacing = Pt(max(1.0, value))
    elif mode == "double":
        paragraph.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph.line_spacing = 2.0
    elif mode == "one_and_half":
        paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.line_spacing = 1.5
    elif mode == "single":
        paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.line_spacing = 1.0
    else:
        paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.line_spacing = max(0.5, value)

    if outline_level is not None:
        p_pr = style.element.get_or_add_pPr()
        _set_or_remove_child(p_pr, "w:outlineLvl", "w:val", str(max(0, min(8, outline_level))))


def _apply_paragraph_style_overrides(style: Any, config: Mapping[str, Any]) -> None:
    """Apply only explicitly supplied properties, preserving template styles."""

    font_keys = {
        "east_asia_font",
        "latin_font",
        "size_pt",
        "bold",
        "italic",
        "underline",
        "color",
        "letter_spacing_pt",
        "font_scale_percent",
    }
    if font_keys.intersection(config):
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        current_latin = style.font.name or "Times New Roman"
        current_east = current_latin
        if r_fonts is not None:
            current_latin = r_fonts.get(qn("w:ascii"), r_fonts.get(qn("w:hAnsi"), current_latin))
            current_east = r_fonts.get(qn("w:eastAsia"), current_east)
        latin = str(config.get("latin_font") or current_latin)
        east_asia = str(config.get("east_asia_font") or current_east)
        if "latin_font" in config:
            style.font.name = latin
        _set_rfonts(r_pr, east_asia, latin)
        if "size_pt" in config:
            style.font.size = Pt(max(1.0, _as_float(config.get("size_pt"), 12.0)))
        if "bold" in config:
            style.font.bold = _as_bool(config.get("bold"), False)
        if "italic" in config:
            style.font.italic = _as_bool(config.get("italic"), False)
        if "underline" in config:
            style.font.underline = _as_bool(config.get("underline"), False)
        if "color" in config:
            color = str(config.get("color", "000000")).strip().lstrip("#")
            if re.fullmatch(r"[0-9A-Fa-f]{6}", color):
                style.font.color.rgb = RGBColor.from_string(color.upper())
        if "letter_spacing_pt" in config:
            spacing = round(_as_float(config.get("letter_spacing_pt"), 0.0) * 20)
            _set_or_remove_child(r_pr, "w:spacing", "w:val", str(spacing))
        if "font_scale_percent" in config:
            scale = max(1, min(600, _as_int(config.get("font_scale_percent"), 100)))
            _set_or_remove_child(r_pr, "w:w", "w:val", str(scale))

    paragraph = style.paragraph_format
    if "alignment" in config:
        alignment = str(config.get("alignment", "left")).strip().lower()
        paragraph.alignment = _ALIGNMENTS.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    if "space_before_pt" in config:
        paragraph.space_before = Pt(max(0.0, _as_float(config.get("space_before_pt"), 0.0)))
    if "space_after_pt" in config:
        paragraph.space_after = Pt(max(0.0, _as_float(config.get("space_after_pt"), 0.0)))
    if "keep_with_next" in config:
        paragraph.keep_with_next = _as_bool(config.get("keep_with_next"), False)
    if "keep_together" in config:
        paragraph.keep_together = _as_bool(config.get("keep_together"), False)
    if "page_break_before" in config:
        paragraph.page_break_before = _as_bool(config.get("page_break_before"), False)
    if "widow_control" in config:
        paragraph.widow_control = _as_bool(config.get("widow_control"), True)
    if "left_indent_cm" in config:
        paragraph.left_indent = Cm(_as_float(config.get("left_indent_cm"), 0.0))
    if "right_indent_cm" in config:
        paragraph.right_indent = Cm(_as_float(config.get("right_indent_cm"), 0.0))
    if "hanging_indent_cm" in config:
        paragraph.first_line_indent = Cm(-abs(_as_float(config.get("hanging_indent_cm"), 0.0)))
    elif "first_line_indent_cm" in config and config.get("first_line_indent_cm") is not None:
        paragraph.first_line_indent = Cm(_as_float(config.get("first_line_indent_cm"), 0.0))
    elif "first_line_indent_chars" in config and config.get("first_line_indent_chars") is not None:
        chars = _as_float(config.get("first_line_indent_chars"), 0.0)
        current_size = style.font.size.pt if style.font.size is not None else 12.0
        size = _as_float(config.get("size_pt"), current_size)
        paragraph.first_line_indent = Pt(chars * size)

    if "line_spacing_mode" in config or "line_spacing_value" in config:
        mode = str(config.get("line_spacing_mode", "multiple")).strip().lower()
        value = _as_float(config.get("line_spacing_value"), 1.0)
        if mode in {"exact", "fixed"}:
            paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.line_spacing = Pt(max(1.0, value))
        elif mode == "at_least":
            paragraph.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph.line_spacing = Pt(max(1.0, value))
        elif mode == "double":
            paragraph.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            paragraph.line_spacing = 2.0
        elif mode == "one_and_half":
            paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph.line_spacing = 1.5
        elif mode == "single":
            paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.line_spacing = 1.0
        else:
            paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.line_spacing = max(0.5, value)


def _get_or_add_style(doc: DocxDocument, name: str, base: str = "Normal"):
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base and style.name != base:
        try:
            style.base_style = doc.styles[base]
        except KeyError:
            pass
    return style


def _configure_styles(doc: DocxDocument, settings: Mapping[str, Any], preserve_template_styles: bool = False) -> None:
    if preserve_template_styles:
        explicit = settings.get("_explicit_settings", {})
        explicit_fonts = explicit.get("fonts", {}) if isinstance(explicit, Mapping) else {}
        global_font_overrides: dict[str, Any] = {}
        if isinstance(explicit_fonts, Mapping):
            if "east_asia" in explicit_fonts:
                global_font_overrides["east_asia_font"] = explicit_fonts["east_asia"]
            if "latin" in explicit_fonts:
                global_font_overrides["latin_font"] = explicit_fonts["latin"]
        body_explicit = explicit.get("body", {}) if isinstance(explicit, Mapping) else {}
        body_overrides = _deep_merge(
            global_font_overrides,
            body_explicit if isinstance(body_explicit, Mapping) else {},
        )
        if isinstance(body_overrides, Mapping) and body_overrides:
            _apply_paragraph_style_overrides(doc.styles["Normal"], body_overrides)

        named_template_bases = {
            "Academic Title": ("title", "Title"),
            "Academic Subtitle": ("subtitle", "Subtitle"),
            "Academic Author": ("author", "Normal"),
            "Academic Metadata": ("metadata", "Normal"),
            "Academic Abstract": ("abstract", "Normal"),
            "Academic Keywords": ("keywords", "Normal"),
            "Academic Reference": ("references", "Normal"),
            "Academic Acknowledgements": ("acknowledgements", "Normal"),
            "Academic Appendix": ("appendix", "Normal"),
            "Academic Header": ("header", "Header"),
            "Academic Footer": ("footer", "Footer"),
        }
        for style_name, (config_name, base) in named_template_bases.items():
            style = _get_or_add_style(doc, style_name, base)
            local_overrides = explicit.get(config_name, {}) if isinstance(explicit, Mapping) else {}
            overrides = _deep_merge(
                global_font_overrides,
                local_overrides if isinstance(local_overrides, Mapping) else {},
            )
            if isinstance(overrides, Mapping) and overrides:
                _apply_paragraph_style_overrides(style, overrides)

        toc_base = "TOC Heading" if "TOC Heading" in [style.name for style in doc.styles] else "Heading 1"
        _get_or_add_style(doc, "Academic TOC Title", toc_base)
        explicit_headings = explicit.get("headings", {}) if isinstance(explicit, Mapping) else {}
        for level in range(1, 6):
            local_overrides = explicit_headings.get(str(level), {}) if isinstance(explicit_headings, Mapping) else {}
            overrides = _deep_merge(
                global_font_overrides,
                local_overrides if isinstance(local_overrides, Mapping) else {},
            )
            if isinstance(overrides, Mapping) and overrides:
                _apply_paragraph_style_overrides(doc.styles[f"Heading {level}"], overrides)
        return

    _apply_paragraph_style(doc.styles["Normal"], settings["body"])
    named = {
        "Academic Title": ("title", "Normal"),
        "Academic Subtitle": ("subtitle", "Normal"),
        "Academic Author": ("author", "Normal"),
        "Academic Metadata": ("metadata", "Normal"),
        "Academic Abstract": ("abstract", "Normal"),
        "Academic Keywords": ("keywords", "Normal"),
        "Academic Reference": ("references", "Normal"),
        "Academic Acknowledgements": ("acknowledgements", "Normal"),
        "Academic Appendix": ("appendix", "Normal"),
        "Academic Header": ("header", "Normal"),
        "Academic Footer": ("footer", "Normal"),
    }
    for style_name, (config_name, base) in named.items():
        style = _get_or_add_style(doc, style_name, base)
        _apply_paragraph_style(style, settings[config_name])

    toc_title_config = _deep_merge(settings["headings"]["1"], {"space_before_pt": 0, "space_after_pt": 18})
    toc_title = _get_or_add_style(doc, "Academic TOC Title", "Normal")
    _apply_paragraph_style(toc_title, toc_title_config)

    for level in range(1, 6):
        heading = doc.styles[f"Heading {level}"]
        _apply_paragraph_style(heading, settings["headings"][str(level)], level - 1)
        q_format = heading.element.find(qn("w:qFormat"))
        if q_format is None:
            heading.element.append(OxmlElement("w:qFormat"))
        toc_style = _get_or_add_style(doc, f"TOC {level}", "Normal")
        toc_config = _deep_merge(
            settings["body"],
            {
                "first_line_indent_chars": 0,
                "left_indent_cm": round((level - 1) * 0.74, 2),
                "space_after_pt": 3,
                "line_spacing_mode": "single",
            },
        )
        _apply_paragraph_style(toc_style, toc_config)


def _configure_page(section: Any, settings: Mapping[str, Any], preserve_template_page: bool = False) -> None:
    explicit = settings.get("_explicit_settings", {})
    explicit_page = explicit.get("page", {}) if isinstance(explicit, Mapping) else {}
    if preserve_template_page and not isinstance(explicit_page, Mapping):
        explicit_page = {}

    width, height, _size, orientation = _page_dimensions(settings)
    size_keys = {"size", "width_mm", "height_mm"}
    if not preserve_template_page or size_keys.intersection(explicit_page):
        section.orientation = WD_ORIENT.LANDSCAPE if orientation == "landscape" else WD_ORIENT.PORTRAIT
        section.page_width = Mm(width)
        section.page_height = Mm(height)
    elif "orientation" in explicit_page:
        current_width = section.page_width.mm
        current_height = section.page_height.mm
        if orientation == "landscape":
            current_width, current_height = max(current_width, current_height), min(current_width, current_height)
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            current_width, current_height = min(current_width, current_height), max(current_width, current_height)
            section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(current_width)
        section.page_height = Mm(current_height)
    margins = settings["page"]["margins_mm"]
    explicit_margins = explicit_page.get("margins_mm", {}) if isinstance(explicit_page, Mapping) else {}
    for name in ("top", "bottom", "left", "right", "gutter"):
        if preserve_template_page and (
            not isinstance(explicit_margins, Mapping) or name not in explicit_margins
        ):
            continue
        value = _as_float(margins.get(name), 0.0)
        if value < 0:
            raise ValueError(f"page.margins_mm.{name} cannot be negative")
        setattr(section, f"{name}_margin" if name != "gutter" else "gutter", Mm(value))
    explicit_header = explicit.get("header", {}) if isinstance(explicit, Mapping) else {}
    explicit_footer = explicit.get("footer", {}) if isinstance(explicit, Mapping) else {}
    if not preserve_template_page or "header_distance_mm" in explicit_page or (
        isinstance(explicit_header, Mapping) and "distance_mm" in explicit_header
    ):
        section.header_distance = Mm(
            _as_float(settings["header"].get("distance_mm"), settings["page"].get("header_distance_mm", 15.0))
        )
    if not preserve_template_page or "footer_distance_mm" in explicit_page or (
        isinstance(explicit_footer, Mapping) and "distance_mm" in explicit_footer
    ):
        section.footer_distance = Mm(
            _as_float(settings["footer"].get("distance_mm"), settings["page"].get("footer_distance_mm", 15.0))
        )


def _set_update_fields_on_open(doc: DocxDocument) -> None:
    settings_element = doc.settings.element
    update = settings_element.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings_element.append(update)
    update.set(qn("w:val"), "true")


def _insert_before_first(parent: Any, element: Any, before_tags: Sequence[str]) -> None:
    tags = {qn(tag) for tag in before_tags}
    for index, child in enumerate(parent):
        if child.tag in tags:
            parent.insert(index, element)
            return
    parent.append(element)


def _configure_page_number_start(section: Any, config: Mapping[str, Any]) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        _insert_before_first(
            sect_pr,
            pg_num_type,
            (
                "w:cols",
                "w:formProt",
                "w:vAlign",
                "w:noEndnote",
                "w:titlePg",
                "w:textDirection",
                "w:bidi",
                "w:rtlGutter",
                "w:docGrid",
                "w:sectPrChange",
            ),
        )
    pg_num_type.set(qn("w:start"), str(max(1, _as_int(config.get("start"), 1))))
    number_format = str(config.get("format", "decimal")).strip()
    format_aliases = {
        "upper_roman": "upperRoman",
        "lower_roman": "lowerRoman",
        "upper_letter": "upperLetter",
        "lower_letter": "lowerLetter",
    }
    pg_num_type.set(qn("w:fmt"), format_aliases.get(number_format.lower(), number_format))


def _clear_story(story: Any) -> None:
    root = story._element
    for child in list(root):
        root.remove(child)
    root.append(OxmlElement("w:p"))


def _set_run_fonts(run: Any, config: Mapping[str, Any]) -> None:
    east_asia = str(config.get("east_asia_font") or "宋体")
    latin = str(config.get("latin_font") or "Times New Roman")
    run.font.name = latin
    run.font.size = Pt(max(1.0, _as_float(config.get("size_pt"), 12.0)))
    run.font.bold = _as_bool(config.get("bold"), False)
    run.font.italic = _as_bool(config.get("italic"), False)
    run.font.underline = _as_bool(config.get("underline"), False)
    color = str(config.get("color", "000000")).strip().lstrip("#")
    if re.fullmatch(r"[0-9A-Fa-f]{6}", color):
        run.font.color.rgb = RGBColor.from_string(color.upper())
    _set_rfonts(run._element.get_or_add_rPr(), east_asia, latin)


def _append_field(paragraph: Any, instruction: str, display_text: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend((begin, instruction_text, separate))
    if display_text:
        display_run = paragraph.add_run(display_text)
        display_run._element.get_or_add_rPr().append(OxmlElement("w:noProof"))
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _add_story_text(story: Any, config: Mapping[str, Any], text: str, style_name: str) -> None:
    paragraph = story.paragraphs[0] if story.paragraphs else story.add_paragraph()
    paragraph.style = style_name
    paragraph.alignment = _ALIGNMENTS.get(str(config.get("alignment", "center")).lower(), WD_ALIGN_PARAGRAPH.CENTER)
    if text:
        run = paragraph.add_run(_clean_markdown_text(text))
        _set_run_fonts(run, config)


def _add_page_number_to_story(story: Any, config: Mapping[str, Any], style_name: str) -> None:
    paragraph = story.add_paragraph() if story.paragraphs[0].text else story.paragraphs[0]
    paragraph.style = style_name
    paragraph.alignment = _ALIGNMENTS.get(str(config.get("alignment", "center")).lower(), WD_ALIGN_PARAGRAPH.CENTER)
    prefix = _clean_markdown_text(config.get("prefix", ""))
    suffix = _clean_markdown_text(config.get("suffix", ""))
    if prefix:
        paragraph.add_run(prefix)
    _append_field(paragraph, " PAGE ", "1")
    if suffix:
        paragraph.add_run(suffix)


def _configure_headers_and_footers(
    doc: DocxDocument,
    paper: Mapping[str, Any],
    settings: Mapping[str, Any],
    preserve_template_stories: bool = False,
) -> None:
    section = doc.sections[0]
    header = settings["header"]
    footer = settings["footer"]
    page_number = settings["page_number"]
    explicit = settings.get("_explicit_settings", {})
    explicit_header = explicit.get("header", {}) if isinstance(explicit, Mapping) else {}
    explicit_footer = explicit.get("footer", {}) if isinstance(explicit, Mapping) else {}
    explicit_page_number = explicit.get("page_number", {}) if isinstance(explicit, Mapping) else {}
    header_requested = isinstance(explicit_header, Mapping) and bool(explicit_header)
    footer_requested = isinstance(explicit_footer, Mapping) and bool(explicit_footer)
    page_number_requested = isinstance(explicit_page_number, Mapping) and bool(explicit_page_number)
    if preserve_template_stories and not (header_requested or footer_requested or page_number_requested):
        return

    number_enabled = _as_bool(page_number.get("enabled"), True) if (not preserve_template_stories or page_number_requested) else False
    number_position = str(page_number.get("position", "footer")).strip().lower()
    if number_position in {"top", "header"}:
        number_position = "header"
    elif number_position in {"bottom", "footer"}:
        number_position = "footer"
    else:
        raise ValueError("page_number.position must be header/top or footer/bottom")

    show_first = _as_bool(page_number.get("show_on_first_page"), False)
    different_first = (
        _as_bool(header.get("different_first_page"), True)
        or _as_bool(footer.get("different_first_page"), True)
        or not show_first
    )
    if not preserve_template_stories or header_requested or footer_requested or page_number_requested:
        section.different_first_page_header_footer = different_first
    if not preserve_template_stories or page_number_requested:
        _configure_page_number_start(section, page_number)

    header_text = str(header.get("text", ""))
    if _as_bool(header.get("use_title"), False) and not header_text:
        header_text = str(paper.get("title", ""))
    footer_text = str(footer.get("text", ""))

    for story, config, text, style_name, position in (
        (section.header, header, header_text, "Academic Header", "header"),
        (section.footer, footer, footer_text, "Academic Footer", "footer"),
    ):
        story_requested = header_requested if position == "header" else footer_requested
        if not preserve_template_stories or story_requested:
            _clear_story(story)
            if _as_bool(config.get("enabled"), False) or text:
                _add_story_text(story, config, text, style_name)
        if number_enabled and number_position == position:
            number_config = _deep_merge(config, page_number)
            _add_page_number_to_story(story, number_config, style_name)

    if different_first:
        for story, config, text, style_name, position in (
            (section.first_page_header, header, header_text, "Academic Header", "header"),
            (section.first_page_footer, footer, footer_text, "Academic Footer", "footer"),
        ):
            story_requested = header_requested if position == "header" else footer_requested
            if not preserve_template_stories or story_requested:
                _clear_story(story)
                if show_first and (_as_bool(config.get("enabled"), False) or text):
                    _add_story_text(story, config, text, style_name)
            if show_first and number_enabled and number_position == position:
                number_config = _deep_merge(config, page_number)
                _add_page_number_to_story(story, number_config, style_name)


def _add_multilevel_numbering(doc: DocxDocument, config: Mapping[str, Any]) -> int | None:
    if not _as_bool(config.get("enabled"), True):
        return None
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        _as_int(element.get(qn("w:abstractNumId")), 0)
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [_as_int(element.get(qn("w:numId")), 0) for element in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id:08X}"[-8:])
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    name = OxmlElement("w:name")
    name.set(qn("w:val"), "AcademicHeadingNumbering")
    abstract.append(name)

    formats = list(config.get("formats") or [])
    number_formats = list(config.get("number_formats") or [])
    starts = list(config.get("starts") or [])
    suffix = str(config.get("suffix", "space")).strip().lower()
    if suffix not in {"tab", "space", "nothing"}:
        suffix = "space"
    for index in range(5):
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), str(index))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), str(max(1, _as_int(starts[index] if index < len(starts) else 1, 1))))
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), str(number_formats[index] if index < len(number_formats) else "decimal"))
        level.append(num_fmt)
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), doc.styles[f"Heading {index + 1}"].style_id)
        level.append(p_style)
        lvl_text = OxmlElement("w:lvlText")
        default_format = ".".join(f"%{level_index}" for level_index in range(1, index + 2))
        lvl_text.set(qn("w:val"), str(formats[index] if index < len(formats) else default_format))
        level.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        level.append(lvl_jc)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), suffix)
        level.append(suff)
        abstract.append(level)

    _insert_before_first(numbering, abstract, ("w:num", "w:numPicBullet"))
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_heading_number(paragraph: Any, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, min(4, level - 1))))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)


def _ensure_page_break(doc: DocxDocument) -> None:
    # Defer the boundary to the next body paragraph. A standalone page-break
    # paragraph can flow onto a new page first and then create an empty page.
    doc._sub2api_page_break_pending = True


def _add_paragraph(doc: DocxDocument, *args: Any, **kwargs: Any) -> Any:
    paragraph = doc.add_paragraph(*args, **kwargs)
    if getattr(doc, "_sub2api_page_break_pending", False):
        paragraph.paragraph_format.page_break_before = True
        doc._sub2api_page_break_pending = False
    return paragraph


def _inline_segments(text: str):
    pattern = re.compile(r"(\*\*.+?\*\*|__.+?__|(?<!\*)\*[^*\n]+?\*(?!\*)|(?<!_)_[^_\n]+?_(?!_))")
    for segment in pattern.split(text):
        if not segment:
            continue
        if (segment.startswith("**") and segment.endswith("**")) or (
            segment.startswith("__") and segment.endswith("__")
        ):
            yield segment[2:-2], True, False
        elif (segment.startswith("*") and segment.endswith("*")) or (
            segment.startswith("_") and segment.endswith("_")
        ):
            yield segment[1:-1], False, True
        else:
            yield segment.replace("`", "").replace("**", "").replace("__", ""), False, False


def _add_formatted_text(paragraph: Any, text: str) -> None:
    for value, bold, italic in _inline_segments(text):
        if not value:
            continue
        run = paragraph.add_run(value)
        if bold:
            run.bold = True
        if italic:
            run.italic = True


def _add_content_paragraphs(doc: DocxDocument, content: Any, style_name: str) -> int:
    count = 0
    for text in _paragraph_blocks(content):
        paragraph = _add_paragraph(doc, style=style_name)
        _add_formatted_text(paragraph, text)
        count += 1
    return count


def _add_heading(
    doc: DocxDocument,
    text: Any,
    level: int,
    settings: Mapping[str, Any],
    num_id: int | None,
    numbered: bool,
) -> Any:
    level = max(1, min(5, level))
    strip_number = numbered and _as_bool(settings["heading_numbering"].get("strip_existing_numbers"), True)
    heading_text = _clean_heading_text(text, strip_number)
    paragraph = _add_paragraph(doc, style=f"Heading {level}")
    _add_formatted_text(paragraph, heading_text)
    if numbered and num_id is not None:
        _apply_heading_number(paragraph, num_id, level)
    return paragraph


def _add_cover(doc: DocxDocument, paper: Mapping[str, Any], settings: Mapping[str, Any]) -> None:
    include_title_page = _as_bool(settings["pagination"].get("include_title_page"), True)
    if include_title_page:
        spacer = _add_paragraph(doc)
        spacer.paragraph_format.space_after = Pt(72)
    title = _add_paragraph(doc, style="Academic Title")
    _add_formatted_text(title, _clean_markdown_text(paper.get("title", "")))
    subtitle_text = _clean_markdown_text(paper.get("subtitle", ""))
    if subtitle_text:
        subtitle = _add_paragraph(doc, style="Academic Subtitle")
        _add_formatted_text(subtitle, subtitle_text)

    author_text = _value_text(paper.get("author"), "、")
    if author_text:
        author = _add_paragraph(doc, style="Academic Author")
        _add_formatted_text(author, author_text)

    metadata_fields = (
        ("institution", "学校"),
        ("department", "院系"),
        ("major", "专业"),
        ("student_id", "学号"),
        ("advisor", "指导教师"),
        ("date", "日期"),
    )
    for key, label in metadata_fields:
        value = _value_text(paper.get(key))
        if value:
            paragraph = _add_paragraph(doc, style="Academic Metadata")
            _add_formatted_text(paragraph, f"{label}：{value}")
    metadata = paper.get("metadata")
    if isinstance(metadata, Mapping):
        metadata = [{"label": key, "value": value} for key, value in metadata.items()]
    if isinstance(metadata, Sequence) and not isinstance(metadata, (str, bytes, bytearray)):
        for item in metadata:
            if isinstance(item, Mapping):
                label = _clean_markdown_text(item.get("label", item.get("name", "")))
                value = _value_text(item.get("value", item.get("text", "")))
                text = f"{label}：{value}" if label else value
            else:
                text = _clean_markdown_text(item)
            if text:
                paragraph = _add_paragraph(doc, style="Academic Metadata")
                _add_formatted_text(paragraph, text)

    if include_title_page and _as_bool(settings["pagination"].get("title_page_break_after"), True):
        _ensure_page_break(doc)


def _add_abstract(doc: DocxDocument, paper: Mapping[str, Any], settings: Mapping[str, Any], num_id: int | None) -> bool:
    abstract = paper.get("abstract")
    keywords = paper.get("keywords")
    if not abstract and not keywords:
        return False
    if abstract:
        _add_heading(doc, settings["labels"].get("abstract", "摘要"), 1, settings, num_id, False)
        _add_content_paragraphs(doc, abstract, "Academic Abstract")
    keyword_text = _value_text(keywords, "；")
    if keyword_text:
        paragraph = _add_paragraph(doc, style="Academic Keywords")
        label = _clean_markdown_text(settings["labels"].get("keywords", "关键词"))
        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        _add_formatted_text(paragraph, keyword_text)
    return True


def _cached_toc_text(paper: Mapping[str, Any], levels: int) -> str:
    lines: list[str] = []
    for section in _all_sections(paper.get("sections", paper.get("outline", []))):
        level = max(1, min(5, _as_int(section.get("level"), 1)))
        if level > levels:
            continue
        title = _clean_heading_text(section.get("title", ""), strip_number=False)
        if title:
            lines.append(f"{'    ' * (level - 1)}{title}")
    return "\n".join(lines) or "目录将在打开文档时自动更新。"


def _add_toc(doc: DocxDocument, paper: Mapping[str, Any], settings: Mapping[str, Any]) -> bool:
    toc = settings["toc"]
    if not _as_bool(toc.get("enabled"), True):
        return False
    if _as_bool(toc.get("page_break_before"), True):
        _ensure_page_break(doc)
    title = _add_paragraph(doc, style="Academic TOC Title")
    _add_formatted_text(title, _clean_markdown_text(toc.get("title", "目录")))
    levels = max(1, min(5, _as_int(toc.get("levels"), 5)))
    field = _add_paragraph(doc, style="TOC 1")
    field.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _append_field(field, f' TOC \\o "1-{levels}" \\h \\z \\u ', _cached_toc_text(paper, levels))
    if _as_bool(toc.get("page_break_after"), True) or _as_bool(
        settings["pagination"].get("toc_page_break_after"), True
    ):
        _ensure_page_break(doc)
    return True


def _render_sections(doc: DocxDocument, paper: Mapping[str, Any], settings: Mapping[str, Any], num_id: int | None) -> int:
    count = 0
    first_level_one = True
    for section in _all_sections(paper.get("sections", paper.get("outline", []))):
        level = section["level"]
        heading_config = settings["headings"][str(level)]
        page_break_before = level == 1 and not first_level_one and (
            _as_bool(settings["pagination"].get("chapter_page_break_before"), True)
            or _as_bool(heading_config.get("page_break_before"), False)
        )
        heading = _add_heading(doc, section["title"], level, settings, num_id, True)
        if page_break_before:
            heading.paragraph_format.page_break_before = True
        _add_content_paragraphs(doc, section["content"], "Normal")
        if level == 1:
            first_level_one = False
        count += 1
    return count


def _add_references(doc: DocxDocument, paper: Mapping[str, Any], settings: Mapping[str, Any], num_id: int | None) -> bool:
    references = paper.get("references")
    if not references:
        return False
    heading = _add_heading(doc, settings["labels"].get("references", "参考文献"), 1, settings, num_id, False)
    if _as_bool(settings["pagination"].get("references_page_break_before"), True):
        heading.paragraph_format.page_break_before = True
    numeric_style = _uses_numeric_reference_style(settings)
    if isinstance(references, Sequence) and not isinstance(references, (str, bytes, bytearray)):
        for index, reference in enumerate(references, start=1):
            text = _numbered_reference_text(reference, index) if numeric_style else _text_from_reference(reference)
            if text:
                paragraph = _add_paragraph(doc, style="Academic Reference")
                _add_formatted_text(paragraph, text)
    elif numeric_style:
        for index, reference in enumerate(_paragraph_blocks(references), start=1):
            text = _numbered_reference_text(reference, index)
            if text:
                paragraph = _add_paragraph(doc, style="Academic Reference")
                _add_formatted_text(paragraph, text)
    else:
        _add_content_paragraphs(doc, references, "Academic Reference")
    return True


def _add_acknowledgements(
    doc: DocxDocument, paper: Mapping[str, Any], settings: Mapping[str, Any], num_id: int | None
) -> bool:
    content = paper.get("acknowledgements", paper.get("thanks"))
    if not content:
        return False
    heading = _add_heading(doc, settings["labels"].get("acknowledgements", "致谢"), 1, settings, num_id, False)
    if _as_bool(settings["pagination"].get("acknowledgements_page_break_before"), True):
        heading.paragraph_format.page_break_before = True
    _add_content_paragraphs(doc, content, "Academic Acknowledgements")
    return True


def _add_appendices(doc: DocxDocument, paper: Mapping[str, Any], settings: Mapping[str, Any], num_id: int | None) -> int:
    appendices = paper.get("appendices", paper.get("appendix", []))
    if not appendices:
        return 0
    if isinstance(appendices, Mapping) or isinstance(appendices, str):
        appendices = [appendices]
    if not isinstance(appendices, Sequence):
        return 0
    count = 0
    for index, appendix in enumerate(appendices, start=1):
        if isinstance(appendix, Mapping):
            title = appendix.get("title") or f"{settings['labels'].get('appendix', '附录')} {index}"
            content = appendix.get("content", appendix.get("text", ""))
        else:
            title = f"{settings['labels'].get('appendix', '附录')} {index}"
            content = appendix
        heading = _add_heading(doc, title, 1, settings, num_id, False)
        if _as_bool(settings["pagination"].get("appendix_page_break_before"), True):
            heading.paragraph_format.page_break_before = True
        _add_content_paragraphs(doc, content, "Academic Appendix")
        count += 1
    return count


def _new_document(template_bytes: bytes | None) -> DocxDocument:
    if template_bytes is None:
        return Document()
    if not isinstance(template_bytes, (bytes, bytearray)) or not template_bytes:
        raise ValueError("template_bytes must contain a valid DOCX file")
    try:
        doc = Document(io.BytesIO(bytes(template_bytes)))
    except Exception as exc:
        raise ValueError("template_bytes must contain a valid DOCX file") from exc
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return doc


def build_academic_paper_docx(
    paper: Mapping[str, Any],
    settings: Mapping[str, Any] | None = None,
    template_bytes: bytes | None = None,
) -> bytes:
    """Render a complete academic paper as DOCX bytes.

    ``paper`` contains structured content; ``settings`` may use either the
    documented nested schema or the administrator form's flat field names.
    When ``template_bytes`` is supplied, its styles, relationships, and final
    section are retained as the starting template before explicit settings are
    applied.
    """

    if not isinstance(paper, Mapping):
        raise TypeError("paper must be a mapping")
    title = _clean_markdown_text(paper.get("title", ""))
    if not title:
        raise ValueError("paper.title is required")

    resolved = _normalise_settings(settings)
    preserve_template = _as_bool(resolved.get("_preserve_template_styles"), False)
    if preserve_template and template_bytes is None:
        raise ValueError("format_preset='uploaded_template' requires template_bytes")
    doc = _new_document(template_bytes)
    _configure_page(doc.sections[0], resolved, preserve_template_page=preserve_template)
    _configure_styles(doc, resolved, preserve_template_styles=preserve_template)
    _set_update_fields_on_open(doc)
    explicit = resolved.get("_explicit_settings", {})
    explicit_numbering = explicit.get("heading_numbering", {}) if isinstance(explicit, Mapping) else {}
    num_id = (
        None
        if preserve_template and not explicit_numbering
        else _add_multilevel_numbering(doc, resolved["heading_numbering"])
    )

    doc.core_properties.title = title
    author = _value_text(paper.get("author"), ", ")
    if author:
        doc.core_properties.author = author

    _add_cover(doc, paper, resolved)
    toc_before_abstract = _as_bool(resolved["toc"].get("before_abstract"), False)
    if toc_before_abstract:
        _add_toc(doc, paper, resolved)
    has_abstract = _add_abstract(doc, paper, resolved, num_id)
    if has_abstract and _as_bool(resolved["pagination"].get("abstract_page_break_after"), True):
        _ensure_page_break(doc)
    if not toc_before_abstract:
        _add_toc(doc, paper, resolved)
    _render_sections(doc, paper, resolved, num_id)
    _add_references(doc, paper, resolved, num_id)
    _add_acknowledgements(doc, paper, resolved, num_id)
    _add_appendices(doc, paper, resolved, num_id)
    _configure_headers_and_footers(doc, paper, resolved, preserve_template_stories=preserve_template)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


__all__ = [
    "ACADEMIC_FORMAT_PRESETS",
    "DEFAULT_ACADEMIC_PAPER_SETTINGS",
    "PAGE_SIZES_MM",
    "build_academic_paper_docx",
    "count_academic_paper_words",
    "summarize_academic_paper",
    "summarize_academic_paper_format",
]
